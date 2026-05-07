from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ─── BRAND COLOURS ────────────────────────────────────────────────────────────
NAVY   = RGBColor(0x0D, 0x2B, 0x55)   # primary navy
GOLD   = RGBColor(0xC9, 0xA0, 0x2D)   # gold accent
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT  = RGBColor(0xF4, 0xF6, 0xF9)   # light grey bg
MID    = RGBColor(0xD0, 0xD8, 0xE4)   # table border grey
TEXT   = RGBColor(0x1A, 0x1A, 0x2E)   # body text
RED_C  = RGBColor(0xC0, 0x39, 0x2B)   # critical red
AMB    = RGBColor(0xD6, 0x8A, 0x00)   # amber / warning
TEAL   = RGBColor(0x0D, 0x6E, 0x6E)   # teal / ok

def set_cell_bg(cell, hex_color: str):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), hex_color)
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_border(cell, color="D0D8E4"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for side in ['top','left','bottom','right']:
        b = OxmlElement(f'w:{"top" if side=="top" else side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), color)
        tcPr.append(b)

def para_fmt(para, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=6):
    para.alignment = align
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after  = Pt(space_after)

def add_run(para, text, bold=False, italic=False, size=10, color=TEXT, font="Calibri"):
    run = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size  = Pt(size)
    run.font.color.rgb = color
    run.font.name  = font
    return run

# ─── DOCUMENT SETUP ───────────────────────────────────────────────────────────
doc = Document()

# Margins
for section in doc.sections:
    section.top_margin    = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin   = Cm(2.2)
    section.right_margin  = Cm(2.2)

# ─── COVER HEADER ─────────────────────────────────────────────────────────────
# Navy banner paragraph
banner = doc.add_paragraph()
banner.paragraph_format.space_before = Pt(0)
banner.paragraph_format.space_after  = Pt(0)
# use shading on the whole paragraph via XML
pPr = banner._p.get_or_add_pPr()
shd = OxmlElement('w:shd')
shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), '0D2B55')
pPr.append(shd)
r1 = banner.add_run("  REGENCARE.IN  ·  DIGITAL GROWTH STRATEGY")
r1.bold = True; r1.font.size = Pt(11); r1.font.color.rgb = GOLD; r1.font.name = "Calibri"

doc.add_paragraph()  # spacer

# Main title
t = doc.add_paragraph()
para_fmt(t, WD_ALIGN_PARAGRAPH.LEFT, 0, 2)
add_run(t, "DIGITAL SALES & MARKETING\n", bold=True, size=22, color=NAVY)
add_run(t, "GO-TO-MARKET (GTM) MASTER LIST", bold=True, size=18, color=GOLD)

sub = doc.add_paragraph()
para_fmt(sub, WD_ALIGN_PARAGRAPH.LEFT, 0, 14)
add_run(sub, "Sequential Action Pipeline — Keyword-Data Driven  ·  Regencare.in  ·  Prepared by NT Global Digital", italic=True, size=10, color=TEAL)

# ─── DOCUMENT PURPOSE ─────────────────────────────────────────────────────────
def section_heading(doc, text):
    h = doc.add_paragraph()
    para_fmt(h, WD_ALIGN_PARAGRAPH.LEFT, 12, 2)
    pPr = h._p.get_or_add_pPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),'0D2B55')
    pPr.append(shd)
    add_run(h, f"  {text}", bold=True, size=12, color=WHITE)

def body_para(doc, text, space_after=5):
    p = doc.add_paragraph()
    para_fmt(p, WD_ALIGN_PARAGRAPH.LEFT, 0, space_after)
    add_run(p, text, size=10, color=TEXT)

def bullet(doc, text, sub=False):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.left_indent  = Inches(0.4 if not sub else 0.7)
    add_run(p, text, size=10, color=TEXT)

def step_heading(doc, step_id, step_name, tagline):
    doc.add_paragraph()
    h = doc.add_paragraph()
    para_fmt(h, WD_ALIGN_PARAGRAPH.LEFT, 6, 0)
    add_run(h, f"STEP {step_id}: {step_name}", bold=True, size=13, color=NAVY)
    s = doc.add_paragraph()
    para_fmt(s, WD_ALIGN_PARAGRAPH.LEFT, 0, 6)
    add_run(s, tagline, italic=True, size=10, color=GOLD)

def sub_step_heading(doc, code, name):
    h = doc.add_paragraph()
    para_fmt(h, WD_ALIGN_PARAGRAPH.LEFT, 8, 2)
    add_run(h, f"{code}. {name}", bold=True, size=11, color=NAVY)

# PURPOSE
section_heading(doc, "Document Purpose")
doc.add_paragraph()
body_para(doc, (
    "This is a flat, sequential action list built specifically for Regencare.in — South India's first dedicated "
    "regenerative medicine centre with branches in Kochi, Calicut, and Chennai. It enumerates every executable "
    "action derived from the keyword baseline dataset already gathered (142 keywords across 6 location and "
    "category segments), ordered chronologically across five stages so the full pipeline is visible at a glance. "
    "Each action is directly mapped to an identified gap, technical issue, or keyword opportunity in the current "
    "Regencare digital footprint. Once reviewed, the next immediate focus area will be assigned from this list."
))

# ─── PIPELINE GLANCE TABLE ────────────────────────────────────────────────────
doc.add_paragraph()
section_heading(doc, "Pipeline at a Glance")
doc.add_paragraph()

tbl = doc.add_table(rows=6, cols=3)
tbl.style = 'Table Grid'
tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

headers = ["Step", "Stage", "Role in Pipeline"]
hdr_fills = ["0D2B55","0D2B55","0D2B55"]
rows_data = [
    ("1", "Data & Automation Foundation", "The Baseline — sets up the data layer and automation rails"),
    ("2", "Core Web & Keyword Strategy",  "The Engine — converts keyword data into a strategic targeting map"),
    ("3", "Content Execution",            "The Build-Out — turns the matrix into live web and editorial assets"),
    ("4", "Off-Page & Multi-Channel Amplification", "The Reach — pushes assets beyond owned properties"),
    ("5", "Voice Search, GEO & AEO Pipelines", "The Next Surface — adapts the playbook for AI and voice answers"),
]
row_fills = ["F4F6F9","FFFFFF","F4F6F9","FFFFFF","F4F6F9"]

# header row
for i, (cell, hdr) in enumerate(zip(tbl.rows[0].cells, headers)):
    set_cell_bg(cell, "0D2B55")
    p = cell.paragraphs[0]
    p.clear()
    para_fmt(p, WD_ALIGN_PARAGRAPH.CENTER, 3, 3)
    add_run(p, hdr, bold=True, size=10, color=WHITE)

for r_idx, (row_data, fill) in enumerate(zip(rows_data, row_fills)):
    row = tbl.rows[r_idx + 1]
    for c_idx, (cell, val) in enumerate(zip(row.cells, row_data)):
        set_cell_bg(cell, fill)
        p = cell.paragraphs[0]
        p.clear()
        align = WD_ALIGN_PARAGRAPH.CENTER if c_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
        para_fmt(p, align, 3, 3)
        bold = c_idx == 0
        col = NAVY if c_idx == 1 else TEXT
        add_run(p, val, bold=bold, size=10, color=col)

# col widths
tbl.columns[0].width = Cm(1.5)
tbl.columns[1].width = Cm(6.5)
tbl.columns[2].width = Cm(9.0)

# ─── IDENTIFIED CRITICAL ISSUES BOX ──────────────────────────────────────────
doc.add_paragraph()
box_h = doc.add_paragraph()
para_fmt(box_h, WD_ALIGN_PARAGRAPH.LEFT, 12, 2)
pPr = box_h._p.get_or_add_pPr()
shd = OxmlElement('w:shd'); shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),'C0392B')
pPr.append(shd)
add_run(box_h, "  CRITICAL ISSUES IDENTIFIED IN BASELINE AUDIT (Immediate Resolution Required)", bold=True, size=10, color=WHITE)

critical_issues = [
    "BRAND CANNIBALIZATION — orthogencare.com (Dr. Vineeth MB's secondary domain) competes directly with regencare.in for the same search queries; this is the single highest-priority SEO threat.",
    "7 H1 TAGS on the homepage — Google's content parsing is confused; one authoritative H1 required.",
    "Blog returns 404 / Resources returns 404 — major content authority signals lost; must be restored before any content production begins.",
    "Booking page is a contact form only — no real booking system; zero conversion tracking; Cal.com integration required.",
    "NAP inconsistency — Chennai branch uses a mismatched email domain; Google Business Profile and Schema data conflict.",
    "GFC Therapy page ranks #9 — 5 dedicated hair clinics rank above; content strategy gap with low-competition opportunity.",
    "Core Web Vitals unverified — PageSpeed API unavailable during audit; manual LCP/CLS/INP check required on mobile and desktop.",
]
for issue in critical_issues:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.left_indent  = Inches(0.4)
    add_run(p, issue, size=10, color=TEXT)

# ══════════════════════════════════════════════════════════════════════════════
# STEP 1
# ══════════════════════════════════════════════════════════════════════════════
step_heading(doc, "1", "Data & Automation Foundation",
    "The Baseline — establish the data layer and automation rails before any market-facing work.")

sub_step_heading(doc, "1A", "Automate Agent Workflows & Processes")
bullets_1a = [
    "Map every recurring Regencare marketing and sales task currently performed manually: lead capture from contact form, appointment confirmations, rank tracking, keyword refresh, Google Business Profile (GBP) updates, and competitor monitoring.",
    "Identify tasks suited for AI agent automation vs. simple scripts: rank-change SERP alerts are scripted; lead routing and content QA are agent-ready.",
    "Agent workflow for keyword rank monitoring — 142-keyword baseline dataset (Kochi, Calicut, Chennai, Kerala Regional, UAE, Conversational) tracked weekly; alert on position changes ≥3 places.",
    "Competitor monitoring agent: daily crawl of orthogencare.com, top-5 GFC hair clinics, and Chennai spine/ortho competitors — flag new pages, new backlinks, or ranking surges.",
    "Lead routing automation: form submission → MSG91 SMS acknowledgement to patient → internal Slack/email alert to clinic team → CRM entry → follow-up trigger at 24 h.",
    "Cal.com booking integration agent: appointment created → MSG91 confirmation → 24-hour reminder → post-appointment feedback request via WhatsApp.",
    "Standard Operating Procedures (SOPs) for each automated workflow so the Regencare team can audit, pause, or extend without developer dependency.",
    "Logging, error handling, and kill switches for every automated workflow before it goes live — no silent failures.",
]
for b in bullets_1a: bullet(doc, b)

sub_step_heading(doc, "1B", "Execute Data Scraping & Initial Market Data Collection")
bullets_1b = [
    "Locked, versioned master keyword dataset — the existing 142-keyword baseline Excel report (v1.0, May 2026) is the starting point; tagged by location cluster, intent, and priority (Critical / High / Medium / Low).",
    "SERP dataset per priority keyword: top 10 organic results, featured snippets, People Also Ask boxes, related searches, and local pack entries — captured for all Critical and High-priority keywords.",
    "Competitor on-page dataset: page titles, meta descriptions, H1 structure, content length, schema markup, and internal link patterns for orthogencare.com, top GFC hair clinics, and UAE medical tourism portals.",
    "Market-level signals: search volume, keyword difficulty, CPC, seasonality (Onam, January–March treatment surge, UAE travel windows), and intent classification per keyword.",
    "Customer-language corpus from Google Reviews (Regencare branches), Practo profiles, health forums, and NRI medical tourism communities — tied to the existing keyword clusters.",
    "Local citation audit: verify NAP (Name, Address, Phone) consistency across Google Business Profile, JustDial, Practo, Sulekha, and hospital directories for all three branches — Chennai email mismatch already identified, resolve first.",
    "Queryable master dataset holding all above — not loose CSVs — so downstream steps pull from a single source of truth.",
    "Recurring scrape and rank-check schedule (weekly for Critical keywords, monthly for Medium/Low) so the dataset reflects live market conditions, not a one-off snapshot.",
]
for b in bullets_1b: bullet(doc, b)

# ══════════════════════════════════════════════════════════════════════════════
# STEP 2
# ══════════════════════════════════════════════════════════════════════════════
step_heading(doc, "2", "Core Web & Keyword Strategy",
    "The Engine — convert raw data into a strategic, prioritised targeting map and fix all technical blockers.")

sub_step_heading(doc, "2A", "Conduct Baseline Website & On-Page SEO Evaluations")
bullets_2a = [
    "Technical SEO audit report: crawlability, indexation status, XML sitemap, robots.txt, canonical tags, redirect chains, and hreflang — regencare.in has multiple issues already identified.",
    "RESOLVE — 7 H1 tags on homepage: consolidate to one primary H1 containing the brand's primary keyword cluster (e.g., 'Regenerative Medicine Centre in Kerala'); demote remaining to H2/H3.",
    "RESOLVE — Blog (404) and Resources (404): restore or rebuild these pages immediately; they represent the primary content authority surface and their absence is a direct organic traffic and E-E-A-T loss.",
    "Core Web Vitals audit: LCP, CLS, INP across mobile and desktop; manually verify using PageSpeed Insights since API was unavailable during baseline audit — flag every page failing thresholds.",
    "Per-URL on-page audit: title tag, meta description, H1/H2 structure, image alt text, internal linking, and schema markup for all existing Regencare pages.",
    "Organic performance benchmark: indexed pages, ranking keywords (against the 142-keyword baseline), top-performing URLs, and traffic trend from Google Search Console.",
    "RESOLVE — Brand cannibalization: document every query where orthogencare.com and regencare.in appear on the same SERP page; build a domain consolidation or content differentiation strategy — this is the #1 SEO priority.",
    "Page-level disposition tag (keep / optimise / consolidate / retire) for every existing Regencare page, scored against its target keyword cluster.",
    "Prioritised fix list: technical blockers first (H1, 404s, Core Web Vitals), then high-impact on-page fixes (title tags, schema, internal linking).",
]
for b in bullets_2a: bullet(doc, b)

sub_step_heading(doc, "2B", "Finalise the Keyword Matrix")
dual = doc.add_paragraph()
para_fmt(dual, WD_ALIGN_PARAGRAPH.LEFT, 4, 4)
add_run(dual, "DUAL-PURPOSE NOTE: ", bold=True, size=10, color=NAVY)
add_run(dual, (
    "The Regencare Keyword Matrix is explicitly designed to serve two functions simultaneously. "
    "(1) It drives organic SEO — informing site architecture, content priorities, and on-page targeting across Kochi, Calicut, Chennai, and the UAE/NRI segment. "
    "(2) It directly dictates Paid Ads targeting — Search and Performance Max campaigns — so the moment ad budget is approved, an immediate market-capture playbook is ready with zero rework."
), size=10, color=TEXT, italic=True)

bullets_2b = [
    "Keyword clusters — all 142 baseline keywords grouped by shared search intent and SERP overlap: Treatment clusters (GFC, PRP, Stem Cell, Platelet Therapy), Condition clusters (hair loss, knee pain, back pain, sports injury), Location clusters (Kochi, Calicut, Chennai, Kerala-wide, UAE/NRI), and Conversational/Cost clusters.",
    "Intent classification per keyword: informational ('what is GFC therapy'), navigational ('regencare kochi'), commercial ('best regenerative medicine clinic Kerala'), transactional ('book GFC therapy appointment').",
    "Funnel-stage tag per keyword: TOFU (awareness — condition/symptom queries) / MOFU (consideration — treatment comparisons) / BOFU (decision — clinic, cost, booking queries).",
    "Composite priority index: search volume × intent value × ranking feasibility ÷ competition — already partially scored in baseline report Priority column (Critical / High / Medium / Low).",
    "Cluster-to-page map — each cluster mapped to a single canonical landing page (one cluster, one page — prevents further cannibalization); e.g., 'GFC therapy Kochi' cluster → /treatments/gfc-therapy/kochi.",
    "Paid Ads launch list: high-CPC, high-intent BOFU keywords (cost queries, booking queries, 'best clinic' queries) flagged for immediate activation when ad budget is released.",
    "Negative keyword lists built from the matrix: exclude irrelevant procedure names, competitor brand terms, and non-treatment queries to protect future ad spend.",
    "UAE/NRI keyword segment (18 keywords in baseline) flagged as low-competition, high-value opportunity — separate landing page cluster for medical tourism positioning.",
    "Locked Keyword Matrix v1.0 (May 2026) — stored as the master baseline; next refresh scheduled 90 days post-website launch.",
]
for b in bullets_2b: bullet(doc, b)

sub_step_heading(doc, "2C", "Finalise Competitor Analysis & Gap Identification")
bullets_2c = [
    "Top organic competitor list per cluster — SERP competitors (may differ from perceived business competitors): orthogencare.com (primary cannibalization threat), top-5 GFC hair clinics ranking above Regencare, Kerala spine and ortho clinics, UAE medical tourism portals.",
    "Top paid competitor dossier: Google Ads copy, landing page structure, and offer angles for any clinic currently running ads on Regencare's priority keywords.",
    "Keyword coverage gap report: all queries where competitors rank in positions 1–10 and Regencare is absent — 'Not in Top 10' rows in the baseline report are the starting inventory.",
    "Content gap and white-space map: topics nobody in the SERP fully owns — especially regenerative medicine research content, condition-specific FAQs, and UAE-facing medical tourism guides.",
    "SERP feature gap report: featured snippets, People Also Ask boxes, video carousels, and local pack positions Regencare could capture with targeted schema and content.",
    "Backlink gap list: high-authority medical, wellness, and local Kerala domains linking to competitors but not to regencare.in — feeds Step 4 outreach.",
    "One-page competitive positioning summary per cluster: where Regencare wins (brand recognition, specialist credentials), where it is vulnerable (brand cannibalization, content volume), and where the white space is (UAE segment, research content, conversational queries).",
]
for b in bullets_2c: bullet(doc, b)

# ══════════════════════════════════════════════════════════════════════════════
# STEP 3
# ══════════════════════════════════════════════════════════════════════════════
step_heading(doc, "3", "Content Execution",
    "The Build-Out — convert the keyword matrix into a live website and an ongoing NMC-compliant content engine.")

sub_step_heading(doc, "3A", "Build New Website Architecture & Landing Pages (Next.js 14 + Strapi v5 CMS)")
bullets_3a = [
    "Site information architecture (IA) driven directly by the keyword cluster map — one cluster maps to one hub or landing page; hub-and-spoke / topic cluster model.",
    "Tech stack: Next.js 14 (App Router, SSG/ISR) for the frontend; Strapi v5 headless CMS for content management so the Regencare team can update treatments, doctors, branches, and blog posts without developer dependency.",
    "URL structure and breadcrumb logic: /treatments/{treatment-slug}, /conditions/{condition-slug}, /doctors/{doctor-slug}, /branches/{branch-slug}, /blog/{post-slug} — all SEO-optimised and cluster-mapped.",
    "CMS content types to configure in Strapi: Treatment, Condition, Doctor, Branch, Blog Post, Research Article, Testimonial, Callback Lead, Site Settings — full editorial control for the client team.",
    "Landing page briefs per cluster: target keyword cluster, search intent, SERP features to target, on-page elements, CTA (Book Appointment / Request Callback), conversion goal.",
    "BOFU/transactional landing pages built first: 'Book GFC Therapy Kochi', 'Book PRP Treatment Calicut', 'Regenerative Medicine Consultation Chennai' — these are also the pages Paid Ads will route to.",
    "Commercial-intent comparison and category pages built second: treatment comparisons, condition-to-treatment guides, 'GFC vs PRP' — high commercial intent, high conversion value.",
    "Informational hub (pillar) pages for each treatment and condition cluster — anchoring the blog content from Step 3B and capturing TOFU traffic for long-term organic growth.",
    "Schema markup implementation per page type: MedicalClinic, MedicalProcedure, FAQPage, Person (Doctor profiles), LocalBusiness (per branch), BreadcrumbList — all NMC-compliant.",
    "Real booking integration: Cal.com embedded on every treatment page and the dedicated /book page — not a contact form; connected to MSG91 SMS confirmation and Nodemailer email confirmation.",
    "Conversion tracking and analytics events on every page before launch: Google Analytics 4, Google Search Console, and Call/Form tracking — never launch a page that cannot be measured.",
    "DPDP Act 2023 compliance: cookie consent banner, privacy policy, data retention policy, and explicit consent checkboxes on all lead capture forms.",
]
for b in bullets_3a: bullet(doc, b)

sub_step_heading(doc, "3B", "Deploy Blog & Content Production Pipeline (NMC-Compliant, E-E-A-T-Led)")
bullets_3b = [
    "Content brief template: target keyword cluster, search intent, top-ranking competitor outline, People Also Ask questions, target word count, internal links required, CTA, NMC compliance checklist.",
    "Editorial calendar driven by keyword matrix priority score — highest ROI clusters first (GFC therapy, PRP Kerala, regenerative medicine Kochi), not what is easiest to write.",
    "E-E-A-T authority strategy: all treatment and medical condition content authored or reviewed by Dr. Vineeth MB (MBBS, MS Ortho) and other named Regencare doctors — author bio, credentials, and profile photo on every clinical article.",
    "End-to-end production pipeline: brief → draft → NMC compliance QA → medical accuracy review → SEO QA → editorial QA → visual assets → publish → distribute → measure.",
    "NMC compliance guardrails on every piece: no guaranteed cure claims, no before/after patient images without consent, no comparative superiority claims, disclaimer on all medical content.",
    "Readability and design standards: scannable formatting, short paragraphs, custom treatment diagrams, clinical research citations, embedded media — content must look more authoritative than SERP competitors.",
    "Original differentiators on every post: Dr. Vineeth MB expert quotes, Regencare patient outcomes data (anonymised, DPDP-compliant), custom treatment pathway diagrams, or unique clinical frameworks.",
    "Internal linking discipline: every new post links to the relevant treatment landing page (Step 3A) and to 2–3 sibling posts — no orphan pages.",
    "Priority content queue (based on baseline keyword gaps): GFC therapy deep-dive, PRP for hair loss, stem cell therapy FAQ, knee pain without surgery, cost guides for all major treatments, UAE medical tourism guide for Keralites abroad.",
    "Content refresh cycle: existing pages re-optimised quarterly — not left to decay; especially critical given the blog was 404 and all historical content authority has been lost.",
    "Per-post KPI dashboard: keyword rankings, organic traffic, scroll depth, time on page, conversion to booking — reviewed monthly and fed back into the editorial calendar.",
]
for b in bullets_3b: bullet(doc, b)

# ══════════════════════════════════════════════════════════════════════════════
# STEP 4
# ══════════════════════════════════════════════════════════════════════════════
step_heading(doc, "4", "Off-Page & Multi-Channel Amplification",
    "The Reach — extend every asset built in Step 3 beyond owned properties to build domain authority and referral traffic.")

sub_step_heading(doc, "4A", "Expand Content Distribution to External Platforms & Channels")
bullets_4a = [
    "Audience platform map for Regencare: Google Business Profile (3 branches — Kochi, Calicut, Chennai), Practo, JustDial, Sulekha, Apollo 24/7 directory, Kerala medical tourism portals, NRI health forums.",
    "Google Business Profile optimisation for all three branches: complete service listings, treatment categories, photo updates, Q&A responses, weekly post cadence, and review response — NAP consistency enforced across all three.",
    "Content repurposing playbook: every long-form blog post sliced into GBP posts, Instagram/Facebook health tips, LinkedIn doctor thought-leadership posts, YouTube Shorts, WhatsApp Business broadcast content.",
    "Digital PR / outreach programme: target health journalists at Kerala-based media (Mathrubhumi, Manorama), national health portals (HealthShots, OnlyMyHealth), and NRI community publications.",
    "Guest posting and contributor pipeline: Dr. Vineeth MB contributions to orthopaedic and regenerative medicine journals, health platforms, and Quora/Reddit medical communities.",
    "Medical directory listings: ensure Regencare appears on Practo, Lybrate, DoctorIndia, and Kerala-specific directories with consistent NAP, treatment list, and booking link.",
    "HARO / journalist-source pipeline: register doctors as expert sources for regenerative medicine stories — earns high-authority editorial backlinks on an ongoing basis.",
    "Backlink outreach list from Step 2C gap analysis: target Kerala wellness websites, medical college publications, and NRI community portals that link to competitors but not Regencare.",
    "Off-page placement tracker: referring domain authority, anchor text used, estimated traffic, and downstream bookings per placement.",
]
for b in bullets_4a: bullet(doc, b)

sub_step_heading(doc, "4B", "Execute Video Creation & Distribution Strategies")
bullets_4b = [
    "Video-priority cluster shortlist: keyword clusters with video-heavy SERPs — GFC therapy, PRP procedure walkthrough, knee pain treatment, regenerative medicine explainers — all have video carousels in Google SERPs.",
    "Video format playbook by funnel stage: patient education explainers (TOFU), treatment procedure walkthroughs and comparison videos (MOFU), anonymised patient journey and testimonial-style case discussions (BOFU — NMC-compliant, no guaranteed outcome claims).",
    "YouTube SEO setup: keyword-optimised titles (e.g., 'What is GFC Therapy? | Regencare Kerala'), keyword-rich descriptions, chapter timestamps, tags, end screens linking to booking page, and custom branded thumbnails.",
    "Short-form vertical cuts for YouTube Shorts, Instagram Reels, and Facebook — derived from long-form procedure videos; each clip targets one conversational keyword from the baseline Conversational & Cost sheet.",
    "Embed videos into the relevant treatment landing pages and blog posts (Step 3) for engagement, dwell-time, and SERP video carousel eligibility.",
    "Transcripts and captions on every video — they double as crawlable text content for SEO and improve accessibility.",
    "Doctor video series: Dr. Vineeth MB (and other Regencare doctors) answering the top People Also Ask questions from the keyword matrix — builds E-E-A-T and supports the GEO/AEO strategy in Step 5.",
    "Video performance dashboard: watch time, retention curves, click-through to booking page, and assisted conversions — reviewed monthly.",
]
for b in bullets_4b: bullet(doc, b)

# ══════════════════════════════════════════════════════════════════════════════
# STEP 5
# ══════════════════════════════════════════════════════════════════════════════
step_heading(doc, "5", "Voice Search, GEO & AEO Pipelines",
    "The Next Surface — adapt the proven text/web playbook for voice interfaces and AI-generated answer engines.")

sub_step_heading(doc, "5A", "Clone & Adapt Successful Pipelines for Voice Search")
bullets_5a = [
    "Voice-search candidate keyword set — the 49-keyword Conversational & Cost cluster in the baseline report (Step 3B content queue) is the primary input; these are already in natural-language question format.",
    "Natural-language re-clustering by spoken intent: 'what is GFC therapy', 'how much does PRP treatment cost in Kochi', 'best regenerative medicine doctor near me' — voice queries are longer, more conversational, and more local than typed queries.",
    "Asset adaptation map: which existing landing pages and blog posts (Step 3A/3B) can be restructured for voice vs. which need new standalone FAQ pages.",
    "Local voice optimisation: 'regenerative medicine clinic near me', 'GFC therapy open today Kochi' — requires complete Google Business Profile data, NAP consistency, and local landing pages for all three branches.",
]
for b in bullets_5a: bullet(doc, b)

sub_step_heading(doc, "5B", "Voice Search Optimisation")
bullets_5b = [
    "Restructured target pages with direct answer blocks (40–60 words) at the top of every FAQ and condition page — this is the format that wins featured snippets and voice read-outs.",
    "Expanded schema implementation: FAQPage schema on every FAQ and blog post, HowTo schema on procedure walkthrough pages, Speakable schema to signal voice-readable content to Google.",
    "FAQ hub answering the full conversational keyword set in plain, spoken language — 'How long does GFC therapy take?', 'Is PRP treatment painful?', 'What is the cost of stem cell therapy in Kerala?'",
    "Test actual voice queries against Google Assistant, Siri, and Alexa to validate what gets read aloud — iterate based on what is and is not returned as a voice answer.",
]
for b in bullets_5b: bullet(doc, b)

sub_step_heading(doc, "5C", "GEO — Generative Engine Optimisation (ChatGPT, Gemini, Perplexity)")
bullets_5c = [
    "GEO target: ensure Regencare.in is cited and named when users ask ChatGPT, Gemini, Perplexity, or Google AI Overviews about regenerative medicine in Kerala, GFC therapy in India, or PRP treatment centres in South India.",
    "Structured knowledge signals: Wikipedia-style factual content about Regencare (founding, speciality, branches, credentials), consistent NAP across all indexed web sources, and a Google Knowledge Panel claim.",
    "Citation-bait content strategy: publish original data, clinical outcome summaries, and treatment cost guides that AI models are likely to cite when answering health queries — these become permanent GEO assets.",
    "Monitor AI answer surfaces monthly: manually query ChatGPT, Gemini, and Perplexity for priority keywords; document whether Regencare is mentioned, quoted, or linked — track over time.",
    "Google AI Overviews optimisation: identify which Regencare queries trigger AI Overviews; ensure the target page has a direct answer block, FAQPage schema, and authoritative authorship in the first 100 words.",
]
for b in bullets_5c: bullet(doc, b)

sub_step_heading(doc, "5D", "AEO — Answer Engine Optimisation (Featured Snippets & Google PAA)")
bullets_5d = [
    "Featured snippet capture targets: identify every conversational keyword in the baseline where the current snippet holder has a weak, unstructured answer — these are immediate takeover opportunities.",
    "People Also Ask (PAA) domination: map all PAA questions appearing on Regencare's priority SERPs; create dedicated answer sections (40–60 words) targeting each PAA box.",
    "Answer format library: paragraph answers (definition queries), numbered list answers (procedure steps), table answers (cost comparisons, treatment comparisons) — match format to SERP feature type.",
    "Feed voice query and AI answer performance back into the master keyword matrix so the loop closes — GEO/AEO intelligence continuously improves the core targeting engine in Step 2.",
]
for b in bullets_5d: bullet(doc, b)

# ─── NEXT ACTION ──────────────────────────────────────────────────────────────
doc.add_paragraph()
section_heading(doc, "STEP → : Next Action")
doc.add_paragraph()
body_para(doc, (
    "Please review the full pipeline above and indicate which step or sub-step the team should focus on first. "
    "Once a focus area is confirmed, a detailed execution plan and resourcing breakdown will be produced for that specific step."
))
doc.add_paragraph()

# TTV box
ttv = doc.add_paragraph()
para_fmt(ttv, WD_ALIGN_PARAGRAPH.LEFT, 4, 4)
pPr = ttv._p.get_or_add_pPr()
shd = OxmlElement('w:shd'); shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),'C9A02D')
pPr.append(shd)
add_run(ttv, "  TIME-TO-VALUE (TTV) INDICATOR: ", bold=True, size=10, color=WHITE)
add_run(ttv, (
    "Step 2B (Keyword Matrix finalisation) is the fastest time-to-value unlock — "
    "completing it activates Paid Search campaigns for immediate lead generation while longer-term SEO assets are being built. "
    "Step 2A (Technical fixes: H1, 404 pages, Core Web Vitals, brand cannibalization) must run in parallel — "
    "no content investment delivers full return on a technically broken site."
), size=10, color=TEXT)

doc.add_paragraph()
body_para(doc, (
    "Step 1 (Data & Automation Foundation) and Step 2 (Core Web & Keyword Strategy) are foundational — most parallel work downstream depends on them being in place. "
    "Step 2C (Competitor Gap) is the secondary high-leverage activation point, since it sharpens both organic targeting and paid ad positioning. "
    "Step 3A (New Website Build) and Step 2A (Technical Fixes) can proceed in parallel once the keyword cluster-to-page map from Step 2B is locked."
), space_after=8)

# ─── FOOTER NOTE ──────────────────────────────────────────────────────────────
footer_p = doc.add_paragraph()
para_fmt(footer_p, WD_ALIGN_PARAGRAPH.LEFT, 12, 2)
pPr = footer_p._p.get_or_add_pPr()
shd = OxmlElement('w:shd'); shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),'F4F6F9')
pPr.append(shd)
add_run(footer_p, "  Prepared by NT Global Digital  ·  Client: Regencare.in  ·  Document Version: 1.0  ·  Date: May 2026  ·  Confidential", size=9, color=NAVY, italic=True)

# ─── SAVE ─────────────────────────────────────────────────────────────────────
out = r"c:\project\AI RESEARCH INTELLIGENCE SYSTEM\Regencare_GTM_Master_List.docx"
doc.save(out)
print(f"Saved: {out}")
