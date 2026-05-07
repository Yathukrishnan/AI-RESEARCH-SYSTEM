"""
SEO & GEO Base Study — Complete Guide for Building a High-Ranking Website
Professional Word Document Generator
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── PALETTE ──────────────────────────────────────────────────────────────────
NAVY  = RGBColor(0x0D,0x2B,0x55); GOLD  = RGBColor(0xC9,0xA0,0x2D)
WHITE = RGBColor(0xFF,0xFF,0xFF); TEXT  = RGBColor(0x1A,0x1A,0x2E)
TEAL  = RGBColor(0x0D,0x6E,0x6E); AMB   = RGBColor(0xD6,0x8A,0x00)
GREEN = RGBColor(0x0D,0x6E,0x3A); RED   = RGBColor(0xC0,0x39,0x2B)
GREY  = RGBColor(0x6C,0x75,0x7D); BLUE  = RGBColor(0x1E,0x3A,0x8A)
PURP  = RGBColor(0x6B,0x21,0xA8); MID   = RGBColor(0xD0,0xD8,0xE4)

def shd(el, hex_):
    s=OxmlElement('w:shd');s.set(qn('w:val'),'clear');s.set(qn('w:color'),'auto');s.set(qn('w:fill'),hex_)
    el._p.get_or_add_pPr().append(s)

def run(para, text, bold=False, italic=False, sz=10, color=TEXT, font="Calibri"):
    r=para.add_run(text);r.bold=bold;r.italic=italic
    r.font.size=Pt(sz);r.font.color.rgb=color;r.font.name=font;return r

def fmt(para, align=WD_ALIGN_PARAGRAPH.LEFT, before=0, after=5):
    para.alignment=align
    para.paragraph_format.space_before=Pt(before)
    para.paragraph_format.space_after=Pt(after)

def banner(doc, text, fill="0D2B55", fg=WHITE, sz=13):
    p=doc.add_paragraph();fmt(p,WD_ALIGN_PARAGRAPH.LEFT,10,2);shd(p,fill)
    run(p,f"  {text}",True,False,sz,fg)

def sub_banner(doc, text, fill="0D2B55", fg=WHITE, sz=10):
    p=doc.add_paragraph();fmt(p,WD_ALIGN_PARAGRAPH.LEFT,0,2);shd(p,fill)
    run(p,f"  {text}",True,False,sz,fg)

def body(doc, text, after=5):
    p=doc.add_paragraph();fmt(p,WD_ALIGN_PARAGRAPH.LEFT,0,after);run(p,text,sz=10,color=TEXT)

def bullet(doc, text, level=0, sz=10, color=TEXT):
    p=doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before=Pt(0);p.paragraph_format.space_after=Pt(3)
    p.paragraph_format.left_indent=Inches(0.35+level*0.25)
    run(p,text,sz=sz,color=color)

def numbered(doc, text, sz=10):
    p=doc.add_paragraph(style='List Number')
    p.paragraph_format.space_before=Pt(0);p.paragraph_format.space_after=Pt(3)
    p.paragraph_format.left_indent=Inches(0.35)
    run(p,text,sz=sz,color=TEXT)

def note(doc, label, text, fill="FFF3CD", label_color=AMB):
    p=doc.add_paragraph();fmt(p,WD_ALIGN_PARAGRAPH.LEFT,4,4);shd(p,fill.replace("#",""))
    run(p,f"  {label}  ",True,False,10,label_color)
    run(p,text,False,False,10,TEXT)

def chapter(doc, num, title, tagline):
    doc.add_paragraph()
    p=doc.add_paragraph();fmt(p,WD_ALIGN_PARAGRAPH.LEFT,6,0);shd(p,"0D2B55")
    run(p,f"  {num}  ",True,False,14,GOLD)
    run(p,title,True,False,14,WHITE)
    s=doc.add_paragraph();fmt(s,WD_ALIGN_PARAGRAPH.LEFT,0,6)
    run(s,f"  {tagline}",False,True,10,TEAL)

def section(doc, title, color=NAVY):
    p=doc.add_paragraph();fmt(p,WD_ALIGN_PARAGRAPH.LEFT,8,2)
    run(p,title,True,False,11,color)

def make_table(doc, headers, rows, col_widths, hdr_fill="0D2B55", alt="F4F6F9"):
    tbl=doc.add_table(rows=1+len(rows),cols=len(headers))
    tbl.style='Table Grid';tbl.alignment=WD_TABLE_ALIGNMENT.LEFT
    for i,(cell,h) in enumerate(zip(tbl.rows[0].cells,headers)):
        p2=cell.paragraphs[0];p2.clear();p2.alignment=WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_before=Pt(2);p2.paragraph_format.space_after=Pt(2)
        shd(p2,hdr_fill);run(p2,h,True,False,9,WHITE)
    for ri,row_data in enumerate(rows):
        fill_=alt if ri%2==0 else "FFFFFF"
        for ci,(cell,val) in enumerate(zip(tbl.rows[ri+1].cells,row_data)):
            p2=cell.paragraphs[0];p2.clear()
            p2.alignment=WD_ALIGN_PARAGRAPH.CENTER if ci==0 else WD_ALIGN_PARAGRAPH.LEFT
            p2.paragraph_format.space_before=Pt(2);p2.paragraph_format.space_after=Pt(2)
            shd(p2,fill_);run(p2,str(val),ci==0,False,9,NAVY if ci==0 else TEXT)
    for col,w in zip(tbl.columns,col_widths): col.width=Cm(w)
    return tbl

def tick_table(doc, headers, rows, col_widths):
    tbl=doc.add_table(rows=1+len(rows),cols=len(headers))
    tbl.style='Table Grid';tbl.alignment=WD_TABLE_ALIGNMENT.LEFT
    for cell,h in zip(tbl.rows[0].cells,headers):
        p2=cell.paragraphs[0];p2.clear();p2.alignment=WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_before=Pt(2);p2.paragraph_format.space_after=Pt(2)
        shd(p2,"0D2B55");run(p2,h,True,False,9,WHITE)
    tick_map={"✓":"1A7A4A","✗":"C0392B","~":"D68A00","?":"6C757D"}
    bg_map={"✓":"E8F5EE","✗":"FDECEA","~":"FFF3CD","?":"F0F0F0"}
    for ri,row_data in enumerate(rows):
        bg="F4F6F9" if ri%2==0 else "FFFFFF"
        for ci,(cell,val) in enumerate(zip(tbl.rows[ri+1].cells,row_data)):
            p2=cell.paragraphs[0];p2.clear()
            p2.alignment=WD_ALIGN_PARAGRAPH.CENTER if ci>0 else WD_ALIGN_PARAGRAPH.LEFT
            p2.paragraph_format.space_before=Pt(2);p2.paragraph_format.space_after=Pt(2)
            if ci>0 and val in tick_map:
                shd(p2,bg_map[val]);run(p2,val,True,False,10,RGBColor.from_string(tick_map[val]))
            else:
                shd(p2,bg);run(p2,str(val),ci==0,False,9,NAVY if ci==0 else TEXT)
    for col,w in zip(tbl.columns,col_widths): col.width=Cm(w)

# ═══════════════════════════════════════════════════════════════════════════
doc=Document()
for sec_ in doc.sections:
    sec_.top_margin=Cm(1.8);sec_.bottom_margin=Cm(1.8)
    sec_.left_margin=Cm(2.2);sec_.right_margin=Cm(2.2)

# ─── COVER ──────────────────────────────────────────────────────────────────
p=doc.add_paragraph();fmt(p,WD_ALIGN_PARAGRAPH.LEFT,0,0);shd(p,"0D2B55")
run(p,"  SEO & GEO BASE STUDY  —  REGENCARE.IN  ·  NT Global Digital",True,False,11,GOLD)

doc.add_paragraph()
t=doc.add_paragraph();fmt(t,WD_ALIGN_PARAGRAPH.LEFT,0,2)
run(t,"COMPLETE GUIDE TO BUILDING\n",True,False,22,NAVY)
run(t,"A HIGH-RANKING, AI-VISIBLE WEBSITE",True,False,18,GOLD)

s=doc.add_paragraph();fmt(s,WD_ALIGN_PARAGRAPH.LEFT,0,14)
run(s,"Content · Technical SEO · Backlinks · Schema · Local SEO · GEO · AEO · Measurement  ·  May 2026",False,True,10,TEAL)

banner(doc,"WHAT THIS DOCUMENT COVERS","0D2B55",WHITE,12)
doc.add_paragraph()
body(doc,(
    "This base study defines every dimension of modern SEO and Generative Engine Optimisation (GEO) "
    "that must be planned, built, and maintained for a new website to reach page 1 and stay there. "
    "It is structured as a reference guide — from site architecture through content, technical "
    "implementation, backlink strategy, local SEO, schema markup, and AI-engine visibility — "
    "with Regencare.in used as the working example throughout. "
    "Each section defines what it is, why it matters, and exactly what to do."
),after=10)

# ─── QUICK REFERENCE TABLE ──────────────────────────────────────────────────
banner(doc,"THE 9 PILLARS — AT A GLANCE","0D2B55",WHITE,12)
doc.add_paragraph()
make_table(doc,
    ["#","Pillar","What It Controls","Impact on Rankings"],
    [
        ("1","Site Architecture","URL structure, crawlability, internal authority flow","Foundational — everything else depends on it"),
        ("2","Technical SEO","Speed, indexation, Core Web Vitals, mobile","Prerequisite — broken technical = low ceiling for all other work"),
        ("3","On-Page SEO","Title tags, H1/H2, meta, content structure","Direct ranking signal — highest ROI per hour of work"),
        ("4","Content Strategy","Depth, E-E-A-T, intent match, freshness","Long-term authority — compounds over time"),
        ("5","Schema Markup","Rich results, AI citations, SERP features","Multiplier — same content earns more SERP features"),
        ("6","Backlinks & Off-Page","Domain authority, trust, referral traffic","Trust signal — takes longest to build, hardest to fake"),
        ("7","Local SEO","Google Maps, local pack, GBP, citations","Critical for any business with physical locations"),
        ("8","GEO — AI Engines","ChatGPT / Gemini / Perplexity citations","Emerging — fast-growing traffic surface, low competition now"),
        ("9","Measurement","Track what works, fix what doesn't, compound gains","Required — without measurement, strategy is guesswork"),
    ],
    [1.0,5.0,6.0,6.0]
)

# ═══════════════════════════════════════════════════════════════════════════
# PILLAR 1 — SITE ARCHITECTURE
# ═══════════════════════════════════════════════════════════════════════════
chapter(doc,"PILLAR 1","SITE ARCHITECTURE",
    "The skeleton of the site — how pages are organised, linked, and discovered by Google.")

body(doc,(
    "Site architecture determines how Google's crawlers move through your site, how "
    "authority flows between pages, and how users navigate from entry to conversion. "
    "A well-structured site makes every other SEO effort more effective. A poorly structured "
    "site puts a ceiling on rankings regardless of how good the content is."
))

section(doc,"1.1  Hub-and-Spoke (Topic Cluster) Model")
body(doc,(
    "Every major topic gets a Hub Page (pillar page) that covers the topic broadly. "
    "Supporting Spoke Pages cover subtopics in depth and link back to the hub. "
    "This concentrates authority on the most important pages and signals topical depth to Google."
))
note(doc,"Regencare Example:",
    "Hub: /treatments/gfc-therapy/  ←  Spokes: /treatments/gfc-therapy/kochi/, "
    "/treatments/gfc-therapy/calicut/, /blog/gfc-vs-prp/, /blog/gfc-therapy-cost-kerala/",
    "E0F4F4",TEAL)

section(doc,"1.2  URL Structure Rules")
bullet(doc,"Short, readable, keyword-containing URLs — /treatments/stem-cell-therapy/ not /page?id=4452")
bullet(doc,"Location modifier at the end — /treatments/prp-therapy/kochi/ not /kochi-prp/")
bullet(doc,"No stop words, no dates in treatment URLs (dates make pages feel stale)")
bullet(doc,"Consistent folder structure: /treatments/, /conditions/, /doctors/, /branches/, /blog/")
bullet(doc,"Never change a URL after indexing without a 301 redirect — Google treats the new URL as a new page")

section(doc,"1.3  Site Architecture Map — What Pages to Build First")
make_table(doc,
    ["Priority","Page Type","URL Pattern","Purpose"],
    [
        ("1 — CRITICAL","Homepage","/","Brand + primary service overview. Links to all treatment and branch hubs."),
        ("1 — CRITICAL","Treatment Hub","/treatments/{treatment}/","One page per treatment. 3,000+ words. All location variants link here."),
        ("1 — CRITICAL","Treatment + Location","/treatments/{treatment}/{city}/","Captures location-specific queries. Inherits hub authority."),
        ("1 — CRITICAL","Branch Page","/branches/{city}/","Local SEO anchor. Contains full NAP, map, hours, photos, local schema."),
        ("2 — HIGH","Condition Hub","/conditions/{condition}/","Condition → treatment bridge. Captures symptom-level queries."),
        ("2 — HIGH","Doctor Profile","/doctors/{name}/","E-E-A-T hub. Credentials, authored content, schema Person entity."),
        ("2 — HIGH","Pricing Page","/pricing/","Captures all cost/price queries — highest-converting traffic type."),
        ("2 — HIGH","Blog / Research","/blog/{slug}/","Informational + long-tail keywords. Links into treatment hubs."),
        ("3 — MEDIUM","FAQ Hub","/faqs/","Captures conversational queries. FAQPage schema. Voice answer source."),
        ("3 — MEDIUM","Comparison Page","/treatments/compare/","GFC vs PRP vs Stem Cell — commercial intent, featured snippet opportunity."),
        ("3 — MEDIUM","Medical Tourism/NRI","/medical-tourism/","UAE/NRI segment. International patient guide."),
        ("4 — LATER","Research Articles","/research/","Clinical commentary. Academic citations. E-E-A-T booster."),
    ],
    [2.5,3.5,5.5,6.5]
)

section(doc,"1.4  Internal Linking Rules")
bullet(doc,"Every treatment page must link to: its parent hub, 2-3 sibling treatment pages, the relevant doctor profile, and the booking/contact page")
bullet(doc,"Every blog post links into 1 hub page and 2-3 related posts — no orphan pages ever")
bullet(doc,"Use descriptive anchor text — 'GFC therapy for hair loss' not 'click here'")
bullet(doc,"Link new pages from existing high-authority pages immediately on publish — accelerates indexing")
bullet(doc,"Breadcrumb navigation on every page — Home > Treatments > GFC Therapy > Kochi")

section(doc,"1.5  Crawlability Essentials")
bullet(doc,"XML sitemap at /sitemap.xml — include all treatment, condition, branch, and blog pages; exclude thank-you pages, admin, search results")
bullet(doc,"robots.txt — allow all treatment pages, block /admin/, /thank-you/, /?s= (search results)")
bullet(doc,"Check Google Search Console > Coverage weekly — fix 'Discovered but not indexed' and 'Crawled but not indexed' errors immediately")
bullet(doc,"No redirect chains — direct 301 only; A→B→C chains dilute link equity")
bullet(doc,"Canonical tags on all pages — especially location variants that share similar content")

# ═══════════════════════════════════════════════════════════════════════════
# PILLAR 2 — TECHNICAL SEO
# ═══════════════════════════════════════════════════════════════════════════
chapter(doc,"PILLAR 2","TECHNICAL SEO",
    "The infrastructure that lets Google discover, crawl, render, and rank your pages.")

body(doc,(
    "Technical SEO is the floor — you cannot rank above a certain level with a broken technical foundation "
    "no matter how good your content is. Google's algorithm now incorporates Core Web Vitals as a "
    "confirmed ranking factor. Mobile performance is the primary evaluation environment."
))

section(doc,"2.1  Core Web Vitals — The Three Metrics Google Measures")
make_table(doc,
    ["Metric","What It Measures","Target","How to Achieve It"],
    [
        ("LCP — Largest Contentful Paint","Time for the largest visible element to load","Under 2.5 seconds","Compress hero images (WebP format). Use CDN. Preload LCP image. Remove render-blocking CSS/JS."),
        ("CLS — Cumulative Layout Shift","Visual stability — do elements jump around?","Under 0.1","Set explicit width/height on all images. Reserve space for ads/embeds. Avoid inserting content above existing content."),
        ("INP — Interaction to Next Paint","Response time to user interactions (clicks, taps)","Under 200ms","Minimise JavaScript execution time. Break long tasks. Use requestAnimationFrame for visual updates."),
    ],
    [3.0,5.5,3.0,7.5]
)
note(doc,"Tool:","PageSpeed Insights (pagespeed.web.dev) — test every key page. Check both Mobile AND Desktop. Mobile score is the ranking signal.","E0F4F4",TEAL)

section(doc,"2.2  Mobile-First — Non-Negotiable")
bullet(doc,"Google indexes the mobile version of your site first. Desktop design is secondary.")
bullet(doc,"Tap targets (buttons, links) must be at least 48×48px — 'Book Appointment' buttons must be thumb-friendly")
bullet(doc,"Font size minimum 16px on mobile — Google flags sub-16px text as a usability issue")
bullet(doc,"No horizontal scrolling on any page — content must fit viewport width")
bullet(doc,"Test every page with Google's Mobile-Friendly Test and Lighthouse mobile audit before launch")

section(doc,"2.3  HTTPS & Security")
bullet(doc,"HTTPS is a confirmed Google ranking signal — every page must serve over HTTPS")
bullet(doc,"Mixed content (HTTP resources on HTTPS pages) kills the HTTPS signal — audit with Chrome DevTools")
bullet(doc,"SSL certificate must auto-renew — a lapsed certificate causes immediate ranking drops and browser warnings")

section(doc,"2.4  Page Speed — Beyond Core Web Vitals")
bullet(doc,"Compress all images: WebP format, max 150KB for content images, max 50KB for thumbnails")
bullet(doc,"Lazy load all images below the fold — do NOT lazy load the hero/LCP image")
bullet(doc,"Minify CSS, JavaScript, and HTML — Next.js does this automatically in production build")
bullet(doc,"Eliminate render-blocking resources — defer non-critical JS, inline critical CSS")
bullet(doc,"Use a CDN (Cloudflare free tier or Vercel Edge Network) — reduces TTFB for users across India and UAE")
bullet(doc,"Target: Google PageSpeed score 90+ on mobile, 95+ on desktop")

section(doc,"2.5  Indexation Checks")
bullet(doc,"'site:regencare.in' in Google — confirms what Google has indexed; check monthly")
bullet(doc,"Google Search Console > Pages > Not Indexed — investigate every 'noindex' tag and 'blocked by robots.txt' entry")
bullet(doc,"Never noindex a page you want to rank — check every page's meta robots tag before launch")
bullet(doc,"Duplicate content: use canonical tags to tell Google which version of a page is the primary one")
bullet(doc,"Thin pages (under 300 words) get 'noindexed' by Google's quality systems — every published page needs substantial content")

section(doc,"2.6  Structured Data — Technical Implementation")
bullet(doc,"Use JSON-LD format (Google's recommended format) — add to <head> section of every page")
bullet(doc,"Link all schema entities via @id — MedicalClinic @id links to Doctor @id links to MedicalProcedure @id")
bullet(doc,"Validate with Google's Rich Results Test after every schema change")
bullet(doc,"Schema errors in Google Search Console must be fixed within 48 hours — they suppress rich results")

# ═══════════════════════════════════════════════════════════════════════════
# PILLAR 3 — ON-PAGE SEO
# ═══════════════════════════════════════════════════════════════════════════
chapter(doc,"PILLAR 3","ON-PAGE SEO",
    "Optimising every element on the page that Google reads to determine relevance and quality.")

body(doc,(
    "On-page SEO is the highest ROI work in the first 90 days of a new site. "
    "It directly signals to Google what each page is about, who it is for, and how authoritative it is. "
    "Every element below is a separate ranking signal."
))

section(doc,"3.1  Title Tags — The Single Most Important On-Page Element")
bullet(doc,"Format: [Primary Keyword] in [City] | [Brand Name]  —  e.g., 'GFC Therapy in Kochi | Regencare'")
bullet(doc,"Length: 50-60 characters — longer titles get truncated in SERPs, losing the keyword")
bullet(doc,"Primary keyword must appear in the first 40 characters — Google weights keyword position in title")
bullet(doc,"Every page gets a unique title — duplicate titles tell Google pages are duplicates of each other")
bullet(doc,"Include a benefit or differentiator where possible — 'GFC Therapy in Kochi — Non-Surgical | Regencare'")

section(doc,"3.2  Meta Descriptions — The SERP Ad Copy")
bullet(doc,"Length: 150-160 characters — longer descriptions get truncated, cutting your CTA")
bullet(doc,"Include: primary keyword, city, doctor name where relevant, and a clear CTA ('Book Free Consultation')")
bullet(doc,"Meta descriptions do NOT directly affect rankings — but they affect Click-Through Rate (CTR), which does")
bullet(doc,"Every page needs a unique meta description — Google rewrites them 70% of the time anyway, but having one improves CTR")
note(doc,"Example:",
    "GFC Therapy in Kochi for hair loss and skin rejuvenation. Non-surgical, done in 45 minutes. "
    "Dr. Vineeth MB, MS Ortho. 3 clinics: Kochi, Calicut, Chennai. Book free consultation today.",
    "E0F4F4",TEAL)

section(doc,"3.3  H1 Tag — One Per Page, Always")
bullet(doc,"ONE H1 tag per page — this is absolute. Multiple H1s confuse Google's topic signal.")
bullet(doc,"H1 must contain the primary keyword — 'GFC Therapy in Kochi' not 'Welcome to Our Clinic'")
bullet(doc,"H1 should match or closely match the title tag — they are the same relevance signal")
bullet(doc,"Current Regencare homepage has 7 H1 tags — this is a critical fix before any other SEO work")

section(doc,"3.4  H2-H6 Heading Structure — The Page Outline Google Reads")
bullet(doc,"H2s are section topics — each one targets a secondary keyword or question")
bullet(doc,"H3s are sub-sections under H2s — provide depth signals without diluting the H2 topic")
bullet(doc,"The winning structure for medical/treatment pages (derived from Apollo, Kokilaben, DermaVue):")
body(doc,"H1: [Treatment] in [City]")
body(doc,"  H2: What is [Treatment]?  →  H2: Conditions Treated  →  H2: Who is Eligible?")
body(doc,"  H2: How Does It Work? (numbered steps)  →  H2: Benefits  →  H2: [Treatment] vs. Alternatives")
body(doc,"  H2: Cost & What to Expect  →  H2: Why Choose Regencare?  →  H2: FAQs (10-12 questions)")
bullet(doc,"Include target keywords naturally in H2s — 'How GFC Therapy Works for Hair Loss' not just 'How It Works'")

section(doc,"3.5  Content Length — Why Word Count Matters")
make_table(doc,
    ["Page Type","Minimum Target","Why","Competitive Benchmark"],
    [
        ("Treatment Hub","3,000+ words","Covers all user questions. Signals topical depth to Google.","DermaVue GFC page: 4,900 words. Ranks #1."),
        ("Condition Page","2,500+ words","Must cover: what is it, causes, symptoms, treatments, FAQs.","Chaitanya OA page: 3,500 words with 6 PubMed citations."),
        ("Location Page","1,800+ words","Location-specific content + full treatment info + local signals.","Apollo Spectra location pages: 1,200-1,400 words — brand compensates."),
        ("Blog Post","1,500+ words","Short posts don't rank for competitive keywords.","Manipal blog: 1,200 words — brand DA compensates. For Regencare: 1,500+ required."),
        ("FAQ Page","1,000+ words","10-15 Q&As of 60-80 words each — comprehensive covers more queries.","Kokilaben FAQ: 8 questions averaging 75 words each."),
        ("Pricing Page","800+ words","Cost factors, process, what's included — not just a price list.","No competitor has this page — first-mover advantage available."),
    ],
    [3.5,3.0,5.0,5.5]
)
note(doc,"Key Insight:","Content length alone does not cause ranking. It is a signal of depth. A 3,000-word page that comprehensively answers the user's query outranks a 5,000-word page that repeats itself. Quality of coverage, not word count, is the goal.","FFF3CD",AMB)

section(doc,"3.6  Keyword Usage — Natural, Not Stuffed")
bullet(doc,"Primary keyword: in H1, first paragraph, at least one H2, and naturally 3-5 times in body — no more")
bullet(doc,"Secondary keywords: in H2s and body — these are variations, related terms, and questions")
bullet(doc,"LSI keywords (semantically related terms): use naturally — Google understands topic context, not just exact matches")
bullet(doc,"Keyword density rule: if a sentence sounds unnatural when read aloud, the keyword is forced — remove it")
bullet(doc,"Search for your target keyword on Google and read the top 3 results — the terms they use are the terms you need")

section(doc,"3.7  Image Optimisation")
bullet(doc,"Every image must have an alt text containing the keyword — 'gfc-therapy-kochi-regencare.webp' not 'image1.jpg'")
bullet(doc,"File name must be descriptive and keyword-containing before upload")
bullet(doc,"WebP format for all content images — 30-40% smaller than JPEG, no quality loss")
bullet(doc,"Add structured data (ImageObject schema) to key treatment images")
bullet(doc,"Compress before upload: target under 150KB for content images, under 80KB for thumbnails")

# ═══════════════════════════════════════════════════════════════════════════
# PILLAR 4 — CONTENT STRATEGY
# ═══════════════════════════════════════════════════════════════════════════
chapter(doc,"PILLAR 4","CONTENT STRATEGY",
    "Building the content that Google trusts, users engage with, and AI engines cite.")

body(doc,(
    "Content is the primary substance of SEO. All technical work and backlink building amplify the content — "
    "but without quality content at the centre, they amplify nothing. Google's 2024+ algorithm explicitly "
    "rewards E-E-A-T content and penalises thin, AI-generated, or untrustworthy medical content."
))

section(doc,"4.1  E-E-A-T — Experience, Expertise, Authoritativeness, Trustworthiness")
body(doc,(
    "E-E-A-T is Google's quality framework for evaluating content, especially in 'Your Money or Your Life' "
    "(YMYL) categories — health, medical, finance, legal. Every medical website is YMYL. "
    "Google uses human quality raters guided by the E-E-A-T framework to calibrate its ranking algorithm."
))
make_table(doc,
    ["Signal","What Google Looks For","How to Implement for Regencare"],
    [
        ("Experience","First-hand experience with the topic","Patient case studies, Dr. Vineeth MB's clinical observations, procedure photos from actual sessions"),
        ("Expertise","Formal qualifications and knowledge","Dr. Vineeth MB byline (MBBS, MS Ortho) on every clinical page. Credentials in Person schema. PubMed citations in articles."),
        ("Authoritativeness","Recognition from the industry","Media mentions, journal citations, speaking engagements, Google Knowledge Panel, Wikipedia-style About page"),
        ("Trustworthiness","Accuracy, transparency, safety","Accurate medical claims, NMC-compliant disclaimers, privacy policy, HTTPS, no misleading before/after, clear contact info"),
    ],
    [3.5,5.5,9.0]
)

section(doc,"4.2  Content Types — What to Produce and Why")
make_table(doc,
    ["Content Type","Primary Goal","SEO Function","Regencare Priority"],
    [
        ("Treatment Pages","Convert visitors to bookings","Rank for commercial keywords + capture BOFU traffic","CRITICAL — build first"),
        ("Condition Pages","Capture symptom-level queries","Bridge between TOFU and treatment pages","HIGH — build second"),
        ("Cost / Pricing Pages","Capture highest-intent traffic","Rank for transactional cost queries — no competitor has these","HIGH — easy win"),
        ("Doctor Profile Pages","Build E-E-A-T across the whole domain","Person schema. Authorship signal for all clinical content.","HIGH — unlocks domain-wide E-E-A-T"),
        ("FAQ Pages / Sections","Capture conversational + voice queries","FAQPage schema → People Also Ask boxes → voice answers","HIGH — attach to every treatment page"),
        ("Blog — Educational","Build topical authority and trust","Rank for informational keywords. Internal link into treatment hubs.","MEDIUM — sustaining strategy"),
        ("Research Commentary","Demonstrate clinical expertise","Highest E-E-A-T signal. Earns backlinks from medical directories.","MEDIUM — Dr. Vineeth MB authored"),
        ("Comparison Pages","Capture comparison-stage queries","Featured snippet opportunity. Commercial intent.","MEDIUM — GFC vs PRP vs Stem Cell"),
        ("NRI / Medical Tourism","Capture UAE/international segment","International keywords with near-zero competition","MEDIUM — new revenue stream"),
        ("Case Studies","Social proof + long-tail traffic","Unique content. E-E-A-T signal. NMC-compliant outcomes.","LOW — ongoing, after launch"),
    ],
    [4.0,4.5,5.5,3.5]
)

section(doc,"4.3  Content Quality Checklist — Every Page Before Publishing")
bullet(doc,"✓  Comprehensively answers the primary search query — read the top 3 results and ensure your page covers more")
bullet(doc,"✓  Written by or reviewed by a named, credentialed doctor — byline visible, linked to doctor profile")
bullet(doc,"✓  1-2 academic citations linked (PubMed or peer-reviewed journal) for clinical claims")
bullet(doc,"✓  No guaranteed outcome claims — use 'may help', 'clinical studies show', 'many patients report'")
bullet(doc,"✓  NMC disclaimer: 'This content is for informational purposes only and does not constitute medical advice'")
bullet(doc,"✓  Internal links: 1 hub link + 2 sibling links + 1 doctor profile link + 1 branch/booking link")
bullet(doc,"✓  FAQ section with 8-12 questions in Q&A format — FAQPage schema applied")
bullet(doc,"✓  At least one original element: doctor quote, clinical data, custom diagram, or unique framework")
bullet(doc,"✓  CTA block at top, middle, and bottom of page — never let user scroll without a conversion opportunity")
bullet(doc,"✓  Reading level: Grade 8-10 (tools: Hemingway App) — clinical accuracy + plain English, not one or the other")

section(doc,"4.4  Content Calendar — What to Build When")
make_table(doc,
    ["Month","Content Priority","Rationale"],
    [
        ("Month 1","Treatment hubs: GFC, PRP, Stem Cell — full rewrites","Highest search volume. Regencare already has URLs — rewrite, not rebuild."),
        ("Month 1","Pricing page","Zero competition. High transactional intent. Easiest ranking win available."),
        ("Month 2","Condition pages: Knee pain, Hair loss, Sports injury","Bridges informational queries to treatment pages. Captures TOFU traffic."),
        ("Month 2","Dr. Vineeth MB profile hub","Unlocks E-E-A-T across all treatment pages immediately on publish."),
        ("Month 3","Location pages: Calicut treatment pages","30-40% lower competition than Kochi. Fastest location-specific ranking."),
        ("Month 3","Chennai branch and treatment pages","New market. Low competition. Fixes NAP issue as a side effect."),
        ("Month 4","GFC vs PRP vs Stem Cell comparison page","Featured snippet opportunity. No competitor has it."),
        ("Month 4","UAE / NRI medical tourism page","Near-zero competition. High LTV patient segment."),
        ("Month 5+","Blog: 2 posts/month — educational, research commentary","Long-tail keyword capture. Sustaining organic growth."),
        ("Ongoing","Content refresh — existing pages quarterly","Prevents decay. Google rewards freshness signal."),
    ],
    [2.5,6.0,9.5]
)

section(doc,"4.5  Content Freshness — Why It Matters")
bullet(doc,"Google measures 'freshness' as a relevance signal — especially for medical content where guidelines change")
bullet(doc,"Add a 'Last Reviewed: [Month Year]' date to all clinical pages — update this when content is reviewed")
bullet(doc,"Quarterly content audit: check Search Console data for pages losing impressions → refresh those pages first")
bullet(doc,"Updating a stale page with new sections, updated statistics, and current schema is faster than writing new pages")

# ═══════════════════════════════════════════════════════════════════════════
# PILLAR 5 — SCHEMA MARKUP
# ═══════════════════════════════════════════════════════════════════════════
chapter(doc,"PILLAR 5","SCHEMA MARKUP",
    "Structured data that tells Google (and AI engines) exactly what your content means.")

body(doc,(
    "Schema markup is the translation layer between your content and machines. "
    "It tells Google not just that a page mentions 'GFC therapy' but that it is a MedicalProcedure, "
    "performed by a named Person (doctor), at a MedicalClinic in a specific Location. "
    "This structured clarity earns rich results in SERPs, improves AI citations, and powers voice answers."
))

section(doc,"5.1  Schema Types for a Medical / Clinic Website")
make_table(doc,
    ["Schema Type","Applied To","What It Unlocks","Priority"],
    [
        ("MedicalClinic","Branch pages (/branches/kochi/)","Local pack eligibility, Knowledge Panel, Maps integration","CRITICAL"),
        ("MedicalProcedure","All treatment pages","Rich result for medical procedures. AI engine citation signal.","CRITICAL"),
        ("MedicalTherapy","Stem cell, PRP, GFC therapy pages","Sub-type of MedicalProcedure — more specific = better AI matching","CRITICAL"),
        ("FAQPage","All pages with FAQ sections","FAQ rich results in SERP — shows Q&As directly under listing. Voice answer source.","CRITICAL"),
        ("Person","Doctor profile pages (/doctors/)","Doctor Knowledge Panel. Authorship signal. Linked to all pages they author.","HIGH"),
        ("LocalBusiness","All branch pages","NAP structured data. Hours, phone, address in SERP.","HIGH"),
        ("BreadcrumbList","Every page","Breadcrumb trail in SERP snippet — improves CTR by ~15%.","HIGH"),
        ("Organization","Homepage","Brand entity. Connects site to Google Knowledge Graph.","HIGH"),
        ("Article / BlogPosting","Blog posts","Article rich results. Author attribution. Date signals.","MEDIUM"),
        ("HowTo","Treatment procedure pages","Numbered step rich results — appears above organic results.","MEDIUM"),
        ("AggregateRating","Pages with patient reviews","Star rating in SERP — dramatically improves CTR.","MEDIUM"),
        ("Speakable","FAQ sections, definition paragraphs","Tells Google which text to read aloud for voice answers.","LOW — future"),
    ],
    [4.5,5.5,6.5,3.0]
)

section(doc,"5.2  JSON-LD Implementation Template")
body(doc,"The recommended pattern is a single JSON-LD block per page that links all entities using @id properties. Example for a GFC Therapy Kochi page:")
p=doc.add_paragraph();fmt(p,WD_ALIGN_PARAGRAPH.LEFT,4,4);shd(p,"1E1E2E")
run(p,
    '{\n'
    '  "@context": "https://schema.org",\n'
    '  "@graph": [\n'
    '    { "@type": "MedicalClinic", "@id": "#clinic-kochi",\n'
    '      "name": "Regencare Kochi", "address": {...}, "telephone": "..." },\n'
    '    { "@type": "MedicalProcedure", "name": "GFC Therapy",\n'
    '      "performedBy": { "@id": "#dr-vineeth" },\n'
    '      "availableAt": { "@id": "#clinic-kochi" } },\n'
    '    { "@type": "Person", "@id": "#dr-vineeth",\n'
    '      "name": "Dr. Vineeth MB", "jobTitle": "Regenerative Medicine Specialist",\n'
    '      "qualification": "MBBS, MS Ortho" },\n'
    '    { "@type": "FAQPage", "mainEntity": [...10 Q&A objects...] },\n'
    '    { "@type": "BreadcrumbList", "itemListElement": [...] }\n'
    '  ]\n'
    '}',
    False,False,9,RGBColor(0x9C,0xD4,0xF8)
)

section(doc,"5.3  Schema Testing & Maintenance")
bullet(doc,"Validate every schema block with Google's Rich Results Test (search.google.com/test/rich-results) before deploying")
bullet(doc,"Check Search Console > Enhancements weekly — fix schema errors within 48 hours")
bullet(doc,"Re-test after every CMS update — plugins and template changes can silently break schema")
bullet(doc,"Never duplicate schema — one MedicalClinic entity per branch, one Person entity per doctor")

# ═══════════════════════════════════════════════════════════════════════════
# PILLAR 6 — BACKLINKS & OFF-PAGE SEO
# ═══════════════════════════════════════════════════════════════════════════
chapter(doc,"PILLAR 6","BACKLINKS & OFF-PAGE SEO",
    "Building the external trust signals that tell Google your site deserves to rank.")

body(doc,(
    "Backlinks are votes of confidence from other websites. A link from a high-authority, "
    "relevant website tells Google that authoritative sources trust your content. "
    "They are the hardest ranking factor to acquire and the hardest to fake — which is why "
    "they carry the most weight in competitive SERPs."
))

section(doc,"6.1  What Makes a Good Backlink")
make_table(doc,
    ["Factor","Good","Bad"],
    [
        ("Source Domain Authority","High DA (40+) health, medical, or local site","Low DA, irrelevant, or spammy sites"),
        ("Relevance","Medical, health, Kerala/India local, ortho/regen topic","Unrelated industry — casino, tech, fashion"),
        ("Link Type","Editorial (in-content) link from a real article","Footer links, sidebar widgets, paid link farms"),
        ("Anchor Text","Descriptive: 'regenerative medicine clinic Kerala'","Over-optimised exact match: 'GFC therapy' on every link"),
        ("Placement","Within the main body content of a relevant page","Footer, sidebar, or link list pages"),
        ("Follow vs NoFollow","DoFollow passes authority","NoFollow does not pass authority — still valuable for traffic"),
    ],
    [4.0,6.0,6.0]
)

section(doc,"6.2  Backlink Acquisition Strategies — Ranked by Quality")
numbered(doc,"Medical Directory Listings (Fastest, Immediate): Practo, Lybrate, JustDial, DoctorIndia, Sulekha — all pass DoFollow links or strong citations. Complete profile immediately on launch.")
numbered(doc,"Google Business Profile (Critical for Local): GBP link to website is a high-authority local signal. Fully optimised GBP for all 3 branches on day 1.")
numbered(doc,"Digital PR — Earn Editorial Links: pitch Dr. Vineeth MB as an expert source to Kerala health journalists (Manorama Health, Mathrubhumi), national health portals (HealthShots, OnlyMyHealth). One earned editorial link from DA 50+ = 50 directory links.")
numbered(doc,"Guest Posting: Dr. Vineeth MB contributes articles to orthopaedic journals, Quora Medical Spaces, LinkedIn Articles — each published piece earns a profile backlink.")
numbered(doc,"HARO (Help A Reporter Out) / SourceBottle: register as a medical expert source. When journalists need a quote on regenerative medicine or PRP, Regencare gets cited with a link.")
numbered(doc,"Broken Link Building: find health pages in Kerala linking to a dead resource on the same topic as Regencare's content — email the webmaster offering Regencare's page as a replacement.")
numbered(doc,"Resource Page Link Building: universities, hospitals, and health NGOs maintain 'resource' pages. Getting Regencare listed as a Kerala regenerative medicine resource earns high-DA links.")
numbered(doc,"Unlinked Brand Mentions: search Google for 'Regencare Kochi' — any mention without a link is a conversion opportunity. Email the site asking them to add the link.")

section(doc,"6.3  Anchor Text Strategy — Avoid Over-Optimisation Penalty")
bullet(doc,"Branded anchors (50%): 'Regencare', 'Regencare Kochi', 'regencare.in' — safest, most natural")
bullet(doc,"Generic anchors (25%): 'click here', 'learn more', 'this clinic', 'visit website' — natural profile")
bullet(doc,"Partial match (15%): 'regenerative medicine Kerala', 'GFC therapy clinic' — keyword variation")
bullet(doc,"Exact match (10% max): 'GFC therapy Kochi' — too many exact match anchors triggers Penguin penalty")

section(doc,"6.4  Backlink Monitoring")
bullet(doc,"Monitor with Ahrefs Free, Moz, or Google Search Console > Links monthly")
bullet(doc,"Disavow toxic links: if a spammy site links to you, disavow via Google Search Console to protect domain authority")
bullet(doc,"Track competitor backlinks monthly — new links they earn are opportunities for you to pursue the same source")
bullet(doc,"Target: 5 new quality backlinks per month in months 1-6; 10+/month from month 7 onward")

# ═══════════════════════════════════════════════════════════════════════════
# PILLAR 7 — LOCAL SEO
# ═══════════════════════════════════════════════════════════════════════════
chapter(doc,"PILLAR 7","LOCAL SEO",
    "Dominating Google Maps, the local 3-pack, and location-specific searches for all three branches.")

body(doc,(
    "For any business with physical locations — especially healthcare — local SEO is the highest-converting "
    "traffic source. A user searching 'GFC therapy near me' or 'regenerative medicine Kochi' who sees "
    "Regencare in the local 3-pack and clicks is 5-10x more likely to book than an organic visitor. "
    "Regencare has three branches — each needs its own local SEO strategy."
))

section(doc,"7.1  Google Business Profile (GBP) — The Local SEO Foundation")
body(doc,"One GBP profile per physical branch. Each profile must be fully completed:")
bullet(doc,"Business name: 'Regencare — Kochi' (not just 'Regencare' — Google uses the city for local pack filtering)")
bullet(doc,"Primary category: 'Medical Clinic' — add secondary categories: 'Orthopedic Surgeon', 'Hair Restoration Service'")
bullet(doc,"Complete service list: list every treatment offered — GFC Therapy, PRP Treatment, Stem Cell Therapy etc.")
bullet(doc,"Hours: accurate and current — wrong hours = negative reviews and ranking drops")
bullet(doc,"Photos: minimum 20 photos — clinic interior, exterior, equipment, doctor photos, team")
bullet(doc,"GBP website link: links to the specific branch page (/branches/kochi/) — not the homepage")
bullet(doc,"Weekly GBP posts: treatment spotlight, health tip, announcement — 1 post/week minimum")
bullet(doc,"Q&A section: proactively add 5-10 common questions with answers — Google shows these in Maps")
bullet(doc,"Booking link: connect Cal.com booking URL directly in GBP — users can book without visiting the website")

section(doc,"7.2  NAP Consistency — Name, Address, Phone")
body(doc,(
    "NAP must be identical across every online mention of each branch — Google Business Profile, website, "
    "Practo, JustDial, Sulekha, directories, social media. Even small differences (Street vs St., "
    "+91 vs 0) confuse Google's local algorithm and reduce local pack eligibility."
))
note(doc,"Regencare Critical Fix:","Chennai branch has a mismatched email domain identified in baseline audit. Fix GBP, website footer, Practo profile, and all citations to use the same contact details. Do this before any other Chennai local SEO work.","FDECEA",RED)
bullet(doc,"Audit all citations with BrightLocal or manually search '[Regencare] [city]' — find every listing and correct NAP")
bullet(doc,"Build citations on: Practo, Lybrate, JustDial, Sulekha, IndiaMedicalTourism, Google Maps (via GBP)")

section(doc,"7.3  Review Strategy")
bullet(doc,"Google Reviews are the #1 local pack ranking factor after proximity — volume AND recency both matter")
bullet(doc,"Target: 50+ reviews per branch within 6 months; 4.7+ average rating maintained")
bullet(doc,"Review generation system: post-appointment WhatsApp message (via MSG91) with a direct Google review link")
bullet(doc,"Respond to every review — positive and negative — within 24 hours. Google rewards engagement.")
bullet(doc,"Never offer incentives for reviews — against Google's policy and risks suspension")
bullet(doc,"Negative reviews: respond professionally, offer to resolve offline, never argue publicly")

section(doc,"7.4  Local Landing Pages")
bullet(doc,"Each branch needs a dedicated page: /branches/kochi/, /branches/calicut/, /branches/chennai/")
bullet(doc,"Each treatment needs a location page: /treatments/gfc-therapy/kochi/, /treatments/gfc-therapy/calicut/")
bullet(doc,"Location pages must contain: full address, embedded Google Map, branch-specific phone, hours, photos, LocalBusiness schema")
bullet(doc,"Avoid copying identical content across location pages — Google treats duplicate location pages as thin content")
bullet(doc,"Include 1-2 paragraphs of location-specific content: local landmarks, parking, public transport, patient catchment area")

# ═══════════════════════════════════════════════════════════════════════════
# PILLAR 8 — GEO & AEO
# ═══════════════════════════════════════════════════════════════════════════
chapter(doc,"PILLAR 8","GEO & AEO — AI ENGINE VISIBILITY",
    "Being cited by ChatGPT, Gemini, Perplexity, and Google AI Overviews — the newest and fastest-growing traffic surface.")

body(doc,(
    "Generative Engine Optimisation (GEO) is the emerging discipline of ensuring your brand is "
    "cited when users ask AI engines questions in your topic area. "
    "Answer Engine Optimisation (AEO) focuses on winning featured snippets and People Also Ask boxes — "
    "the original AI-adjacent SERP features. Both are now essential, especially as AI Overview "
    "(Google's AI-generated answer box) appears for an estimated 30-40% of health queries."
))

section(doc,"8.1  How AI Engines Decide What to Cite")
make_table(doc,
    ["AI Engine","Primary Source Data","What Gets Cited","How to Get Cited"],
    [
        ("Google AI Overviews","Google's indexed web + knowledge graph","Pages with strong E-E-A-T, FAQPage schema, and direct answer blocks in first 100 words","FAQ schema + named doctor authorship + direct 40-60 word answer in intro"),
        ("ChatGPT (GPT-4o)","Training data + Bing web search","Frequently mentioned brands on authoritative sites. Consistent entity data.","High DA backlinks + brand mentions on medical directories + consistent NAP across the web"),
        ("Perplexity AI","Live web search + cited sources","Pages that directly and concisely answer the query with clear factual structure","Structured content: headers, bullet points, factual claims, source citations"),
        ("Google Gemini","Google's index + Knowledge Graph","Same as Google AI Overviews — E-E-A-T signals + structured data","Schema markup + Google Knowledge Panel + Wikipedia/Wikidata presence"),
    ],
    [3.5,5.5,5.5,6.0]
)

section(doc,"8.2  GEO Content Strategy — What to Produce")
bullet(doc,"Direct answer blocks: every FAQ answer must be 40-60 words, structured as: '[Question] — [Answer in plain English.]' — this is the exact format AI models extract and repeat")
bullet(doc,"Entity consistency: 'Regencare', 'Regencare.in', 'Regencare Kochi' — same name, same description, same facts across all pages, GBP, Practo, social media, and press mentions")
bullet(doc,"About page — Wikipedia-style: founding date, mission, speciality, doctor names, location count, accreditations — factual, third-person, AI-readable")
bullet(doc,"Citation-bait content: original data that AI models want to cite — 'Cost of GFC therapy in Kerala ranges from ₹X to ₹Y per session' — a specific, factual, citable claim")
bullet(doc,"Brand mentions campaign: get Regencare named on high-DA sites (medical directories, news articles, health portals) — more web mentions = higher AI model confidence in citing the brand")

section(doc,"8.3  AEO — Winning Featured Snippets & People Also Ask")
body(doc,"Featured snippets are the direct answer boxes that appear above organic results. Winning a featured snippet means appearing above even the #1 organic result.")
bullet(doc,"Paragraph snippets (definition queries): 40-60 word direct answer immediately after the H2 heading — 'What is GFC therapy? GFC (Growth Factor Concentrate) therapy is...'")
bullet(doc,"List snippets (how-to, step queries): numbered list with 5-8 items — Google displays the list directly in SERPs")
bullet(doc,"Table snippets (comparison queries): HTML table comparing treatment options — Google displays the table in SERPs")
bullet(doc,"People Also Ask (PAA): every FAQ section question that exactly matches a PAA question found in Google is a PAA box takeover opportunity — research actual PAA boxes for your target keywords")
bullet(doc,"Monitor featured snippets monthly: use Search Console to identify queries where you appear in position 1-3 but don't have the snippet — those pages need a 40-60 word direct answer block added")

section(doc,"8.4  Voice Search Optimisation")
bullet(doc,"Voice queries are longer, conversational, and local: 'Where can I get GFC therapy in Kochi?' not 'GFC therapy Kochi'")
bullet(doc,"Conversational keyword cluster (49 keywords already in Regencare's baseline): answer every one of these in the FAQ hub")
bullet(doc,"Local voice: 'regenerative medicine clinic near me' — requires fully optimised GBP + correct LocalBusiness schema + high review count")
bullet(doc,"Speakable schema: mark FAQ answer sections with Speakable schema — tells Google's voice assistant which text to read aloud")
bullet(doc,"Test manually: ask Google Assistant, Siri, and Alexa your 10 priority voice queries. Document what gets returned. Iterate content until Regencare's answer is returned.")

section(doc,"8.5  Google Knowledge Panel")
bullet(doc,"A Knowledge Panel is Google's information card about your business — appears for branded searches on desktop")
bullet(doc,"Claim your Knowledge Panel via 'Claim this knowledge panel' — requires Google to verify identity")
bullet(doc,"Ensure GBP data, website About page, and all directory listings use identical information — consistency builds the Knowledge Panel")
bullet(doc,"Add social media profiles (Instagram, Facebook, LinkedIn, YouTube) to GBP — Google links them in the Knowledge Panel")

# ═══════════════════════════════════════════════════════════════════════════
# PILLAR 9 — MEASUREMENT
# ═══════════════════════════════════════════════════════════════════════════
chapter(doc,"PILLAR 9","MEASUREMENT & MONITORING",
    "Tracking every signal so you know what's working, what's not, and where to invest next.")

body(doc,(
    "SEO without measurement is guesswork. Every action in the previous eight pillars needs "
    "a corresponding tracking mechanism so results can be attributed, strategies can be refined, "
    "and ROI can be demonstrated. Set up all measurement before launching a single page."
))

section(doc,"9.1  Essential Toolset — Free Stack")
make_table(doc,
    ["Tool","What to Track","Frequency","Action Trigger"],
    [
        ("Google Search Console","Impressions, clicks, CTR, position per keyword. Coverage errors. Core Web Vitals. Schema errors.","Weekly","Any page losing 20%+ impressions week-over-week needs investigation"),
        ("Google Analytics 4","Sessions, users, engagement rate, conversions (form fills, booking clicks, calls), traffic sources.","Weekly","Any traffic channel dropping 15%+ week-over-week needs cause analysis"),
        ("Google Business Profile Insights","Profile views, direction requests, call clicks, website clicks per branch.","Monthly","Low call-click rate = weak CTA or wrong category. Low profile views = need more photos/posts."),
        ("PageSpeed Insights","Core Web Vitals per page — LCP, CLS, INP on mobile.","Monthly or after major site changes","Any page failing Core Web Vitals thresholds needs technical fix before other SEO work"),
        ("Ubersuggest / Ahrefs Free","Keyword rankings, new backlinks, competitor new content.","Monthly","Competitor gains 10+ new backlinks on a keyword = content gap to fill"),
        ("Google Alerts","Brand mentions: 'Regencare', 'Dr. Vineeth MB', competitor alerts.","As they happen","Unlinked mentions = outreach opportunity. Negative mentions = PR response needed"),
    ],
    [4.5,5.5,3.0,6.5]
)

section(doc,"9.2  KPIs — What Success Looks Like Month by Month")
make_table(doc,
    ["Timeframe","KPI Target","Leading Indicators"],
    [
        ("Month 1-2 (Foundation)","Technical fixes complete. GBP fully optimised × 3 branches. GSC showing 0 critical errors.","Search Console coverage 100% valid. PageSpeed mobile 85+. Schema validated."),
        ("Month 3 (Early Traction)","First 5 keywords on page 1. Local pack appearance for 2+ branch locations.","Impressions rising week-over-week. CTR above 3% for branded queries."),
        ("Month 4-6 (Growth)","20+ keywords page 1. Organic traffic 500+ sessions/month. 3 bookings/month from organic.","Average position improving from 15 → 8 across tracked keywords."),
        ("Month 7-12 (Compounding)","50+ keywords page 1. Organic traffic 2,000+ sessions/month. 10+ bookings/month from organic.","Domain Authority growing. New backlinks monthly. Featured snippets won."),
        ("Year 2+","Top 3 for all primary treatment keywords. Local pack top 3 × all 3 branches. AI citations confirmed.","Citation in ChatGPT/Gemini for Kerala regen queries. Knowledge Panel live."),
    ],
    [3.5,6.5,8.0]
)

section(doc,"9.3  Monthly SEO Reporting — What to Include")
bullet(doc,"Keyword rankings: top 20 keywords — position this month vs. last month vs. launch date")
bullet(doc,"Organic traffic: sessions, users, engagement rate — month-over-month and year-over-year comparison")
bullet(doc,"Conversions from organic: booking form submissions, Call Now clicks, WhatsApp initiations")
bullet(doc,"Local pack: GBP profile views, direction requests, call clicks per branch")
bullet(doc,"New backlinks earned this month: domain, DA, anchor text, link type")
bullet(doc,"Content published this month: title, URL, target keyword, word count")
bullet(doc,"Technical health: Core Web Vitals status, crawl errors, schema errors, coverage issues")
bullet(doc,"AI visibility check: manual test of 10 priority queries in ChatGPT, Gemini, Perplexity — is Regencare cited?")

# ─── QUICK REFERENCE CHECKLIST ──────────────────────────────────────────────
doc.add_paragraph()
banner(doc,"MASTER CHECKLIST — BEFORE SITE LAUNCH","0D2B55",WHITE,12)
doc.add_paragraph()
CHECKLIST=[
    ("ARCHITECTURE","Hub-and-spoke URL structure defined and built",True),
    ("ARCHITECTURE","XML sitemap submitted to Google Search Console",True),
    ("ARCHITECTURE","robots.txt configured — no treatment pages blocked",True),
    ("ARCHITECTURE","Breadcrumb navigation on every page",True),
    ("TECHNICAL","Core Web Vitals: mobile LCP <2.5s, CLS <0.1, INP <200ms",True),
    ("TECHNICAL","PageSpeed mobile score 85+",True),
    ("TECHNICAL","HTTPS on all pages — zero mixed content",True),
    ("TECHNICAL","Google Analytics 4 + Search Console connected",True),
    ("TECHNICAL","Conversion tracking on booking form, Call, WhatsApp",True),
    ("ON-PAGE","Single H1 on every page containing primary keyword",True),
    ("ON-PAGE","Unique title tag 50-60 chars on every page",True),
    ("ON-PAGE","Meta description 150-160 chars on every page",True),
    ("ON-PAGE","All images have alt text + WebP format + <150KB",True),
    ("ON-PAGE","Treatment pages: 3,000+ words with 9+ H2 sections",True),
    ("CONTENT","Named doctor (Dr. Vineeth MB) byline on all clinical pages",True),
    ("CONTENT","1-2 PubMed/academic citations on all treatment pages",True),
    ("CONTENT","NMC disclaimer on every medical content page",True),
    ("CONTENT","FAQ section (8-12 Q&As) on every treatment page",True),
    ("CONTENT","Pricing page published before launch",True),
    ("SCHEMA","MedicalClinic schema on all branch pages",True),
    ("SCHEMA","MedicalProcedure schema on all treatment pages",True),
    ("SCHEMA","FAQPage schema on all FAQ sections",True),
    ("SCHEMA","Person schema for Dr. Vineeth MB — linked from all authored pages",True),
    ("SCHEMA","BreadcrumbList schema on every page",True),
    ("SCHEMA","All schema validated in Rich Results Test — zero errors",True),
    ("LOCAL SEO","GBP fully optimised × 3 branches (Kochi, Calicut, Chennai)",True),
    ("LOCAL SEO","NAP consistent across all 3 branches — website, GBP, directories",True),
    ("LOCAL SEO","Chennai email NAP inconsistency fixed",True),
    ("LOCAL SEO","Practo, JustDial, Sulekha listings completed for all branches",True),
    ("LOCAL SEO","Review generation system active (MSG91 WhatsApp post-appointment)",True),
    ("BACKLINKS","Medical directory listings submitted: Practo, Lybrate, JustDial",True),
    ("BACKLINKS","GBP profile link pointing to website — active",True),
    ("GEO / AEO","Direct answer blocks (40-60 words) in first H2 of every page",True),
    ("GEO / AEO","About page published — factual, Wikipedia-style entity content",True),
    ("GEO / AEO","Google Knowledge Panel claimed and verified",True),
    ("GEO / AEO","Voice query test: 10 priority queries tested on Google Assistant",True),
    ("BOOKING","Cal.com booking live — Book Appointment links active on all pages",True),
    ("BOOKING","CTA block appears at top, middle, and bottom of every treatment page",True),
]
for item_group in [CHECKLIST[i:i+5] for i in range(0,len(CHECKLIST),5)]:
    for category,item,_ in item_group:
        p=doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before=Pt(0);p.paragraph_format.space_after=Pt(2)
        p.paragraph_format.left_indent=Inches(0.35)
        run(p,f"[{category}]  ",True,False,9,NAVY)
        run(p,item,False,False,9,TEXT)

# ─── FOOTER ─────────────────────────────────────────────────────────────────
doc.add_paragraph()
fp=doc.add_paragraph();fmt(fp,WD_ALIGN_PARAGRAPH.LEFT,10,2);shd(fp,"F4F6F9")
run(fp,"  Prepared by NT Global Digital  ·  Client: Regencare.in  ·  SEO & GEO Base Study  ·  Version 1.0  ·  May 2026  ·  Confidential",False,True,9,NAVY)

out=r"c:\project\AI RESEARCH INTELLIGENCE SYSTEM\Regencare_SEO_GEO_Base_Study.docx"
doc.save(out)
print(f"Saved: {out}")
