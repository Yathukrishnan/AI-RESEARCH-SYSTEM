from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── COLOURS ──────────────────────────────────────────────────────────────────
NAVY   = RGBColor(0x0D, 0x2B, 0x55)
GOLD   = RGBColor(0xC9, 0xA0, 0x2D)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT  = RGBColor(0xF4, 0xF6, 0xF9)
TEXT   = RGBColor(0x1A, 0x1A, 0x2E)
RED_C  = RGBColor(0xC0, 0x39, 0x2B)
GREEN  = RGBColor(0x0D, 0x6E, 0x3A)
TEAL   = RGBColor(0x0D, 0x6E, 0x6E)
AMB    = RGBColor(0xD6, 0x8A, 0x00)

def set_cell_bg(cell, hex_color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), hex_color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_run(para, text, bold=False, italic=False, size=10, color=TEXT, font="Calibri"):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = font
    return run

def para_fmt(para, align=WD_ALIGN_PARAGRAPH.LEFT, before=0, after=5):
    para.alignment = align
    para.paragraph_format.space_before = Pt(before)
    para.paragraph_format.space_after  = Pt(after)

def section_banner(doc, text, fill="0D2B55", text_color=WHITE):
    h = doc.add_paragraph()
    para_fmt(h, WD_ALIGN_PARAGRAPH.LEFT, 12, 2)
    pPr = h._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), fill)
    pPr.append(shd)
    add_run(h, f"  {text}", bold=True, size=12, color=text_color)

def phase_heading(doc, code, title, tagline):
    doc.add_paragraph()
    h = doc.add_paragraph()
    para_fmt(h, WD_ALIGN_PARAGRAPH.LEFT, 8, 0)
    add_run(h, f"{code}: {title}", bold=True, size=13, color=NAVY)
    s = doc.add_paragraph()
    para_fmt(s, WD_ALIGN_PARAGRAPH.LEFT, 0, 5)
    add_run(s, tagline, italic=True, size=10, color=GOLD)

def sub_heading(doc, text):
    h = doc.add_paragraph()
    para_fmt(h, WD_ALIGN_PARAGRAPH.LEFT, 8, 2)
    add_run(h, text, bold=True, size=11, color=NAVY)

def body(doc, text, after=5):
    p = doc.add_paragraph()
    para_fmt(p, WD_ALIGN_PARAGRAPH.LEFT, 0, after)
    add_run(p, text, size=10, color=TEXT)

def bullet(doc, text, indent=0.4):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.left_indent  = Inches(indent)
    add_run(p, text, size=10, color=TEXT)

def note_box(doc, label, text, fill="F4F6F9", label_color=None):
    label_color = label_color or NAVY
    p = doc.add_paragraph()
    para_fmt(p, WD_ALIGN_PARAGRAPH.LEFT, 4, 4)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), fill)
    pPr.append(shd)
    add_run(p, f"  {label} ", bold=True, size=10, color=label_color)
    add_run(p, text, size=10, color=TEXT)

# ─── HELPER: FULL-WIDTH TABLE ──────────────────────────────────────────────────
def make_table(doc, headers, rows, col_widths, hdr_fill="0D2B55", alt_fill="F4F6F9"):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    # header
    for i, (cell, hdr) in enumerate(zip(tbl.rows[0].cells, headers)):
        set_cell_bg(cell, hdr_fill)
        p = cell.paragraphs[0]; p.clear()
        para_fmt(p, WD_ALIGN_PARAGRAPH.CENTER, 3, 3)
        add_run(p, hdr, bold=True, size=9, color=WHITE)
    # data rows
    for r_idx, row_data in enumerate(rows):
        fill = alt_fill if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, (cell, val) in enumerate(zip(tbl.rows[r_idx + 1].cells, row_data)):
            set_cell_bg(cell, fill)
            p = cell.paragraphs[0]; p.clear()
            align = WD_ALIGN_PARAGRAPH.CENTER if c_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
            para_fmt(p, align, 2, 2)
            bold = c_idx == 0
            add_run(p, str(val), bold=bold, size=9, color=TEXT)
    # col widths
    for col, w in zip(tbl.columns, col_widths):
        col.width = Cm(w)
    return tbl

# ══════════════════════════════════════════════════════════════════════════════
# BUILD DOCUMENT
# ══════════════════════════════════════════════════════════════════════════════
doc = Document()
for section in doc.sections:
    section.top_margin    = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin   = Cm(2.2)
    section.right_margin  = Cm(2.2)

# ─── COVER BANNER ─────────────────────────────────────────────────────────────
banner = doc.add_paragraph()
para_fmt(banner, WD_ALIGN_PARAGRAPH.LEFT, 0, 0)
pPr = banner._p.get_or_add_pPr()
shd = OxmlElement('w:shd'); shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),'0D2B55')
pPr.append(shd)
add_run(banner, "  REGENCARE.IN  ·  COMPETITIVE INTELLIGENCE STUDY", bold=True, size=11, color=GOLD)

doc.add_paragraph()

t = doc.add_paragraph()
para_fmt(t, WD_ALIGN_PARAGRAPH.LEFT, 0, 2)
add_run(t, "COMPETITIVE ANALYSIS\n", bold=True, size=22, color=NAVY)
add_run(t, "STUDY PLAN & EXECUTION FRAMEWORK", bold=True, size=18, color=GOLD)

sub = doc.add_paragraph()
para_fmt(sub, WD_ALIGN_PARAGRAPH.LEFT, 0, 14)
add_run(sub, "Regencare.in  ·  Prepared by NT Global Digital  ·  Version 1.0  ·  May 2026", italic=True, size=10, color=TEAL)

# ─── PURPOSE ──────────────────────────────────────────────────────────────────
section_banner(doc, "Document Purpose")
doc.add_paragraph()
body(doc, (
    "This document defines the complete methodology, phase-by-phase execution plan, toolset, and deliverables "
    "for the Regencare.in competitive intelligence study. The study maps every digital competitor currently "
    "outranking or threatening Regencare.in across organic search, paid advertising, local search, content, "
    "and AI answer engines. All findings feed directly into Steps 2B (Keyword Matrix), 3A (Website Architecture), "
    "3B (Content Pipeline), and 4A (Off-Page Amplification) of the GTM Master List."
), after=10)

# ─── SCOPE TABLE ──────────────────────────────────────────────────────────────
section_banner(doc, "Study Scope at a Glance")
doc.add_paragraph()

make_table(doc,
    headers=["Phase", "Focus Area", "Key Question Answered", "Primary Output"],
    rows=[
        ("P1", "Competitor Identification",    "Who are we actually competing against in SERPs?",          "Confirmed competitor list per cluster"),
        ("P2", "SERP & Ranking Analysis",      "Where do they rank and what features do they own?",        "SERP position map + feature gap list"),
        ("P3", "On-Page & Content Audit",      "What do their pages look like and how deep is the content?","On-page audit table + content gap map"),
        ("P4", "Keyword Gap Analysis",         "What keywords do they rank for that we do not?",           "Keyword gap report"),
        ("P5", "Backlink Gap Analysis",        "Who links to them but not to us?",                         "Backlink outreach hit list"),
        ("P6", "Paid Ads Intelligence",        "What are they spending on and what offers are they making?","Ads dossier per cluster"),
        ("P7", "Local & GBP Intelligence",     "How do they dominate local pack and maps?",                "Local SEO gap report"),
        ("P8", "AI & Answer Engine Audit",     "Who is being cited in ChatGPT, Gemini, Perplexity?",       "GEO/AEO competitor map"),
        ("P9", "White Space Identification",   "What does nobody own yet?",                                "Opportunity matrix"),
        ("P10","Positioning Summary",          "Where do we win, where are we vulnerable?",                "One-page positioning brief per cluster"),
    ],
    col_widths=[1.5, 4.2, 6.8, 4.5],
)

# ─── COMPETITOR TIERS ─────────────────────────────────────────────────────────
doc.add_paragraph()
section_banner(doc, "Competitor Tiers — Who We Are Studying")
doc.add_paragraph()
body(doc, (
    "Competitors are split into three tiers. Tier 1 is the immediate threat — entities already directly "
    "displacing Regencare.in in search results. Tier 2 is the organic SERP field per keyword cluster. "
    "Tier 3 is the paid and AI-surface competitive set."
), after=8)

make_table(doc,
    headers=["Tier", "Competitor Type", "Known Examples", "Priority"],
    rows=[
        ("1", "Brand Cannibalization",        "orthogencare.com (Dr. Vineeth MB's secondary domain — same doctor, competing domain)",           "CRITICAL"),
        ("1", "Direct Treatment Competitors", "Top-5 hair clinics ranking above Regencare for GFC therapy keywords",                            "CRITICAL"),
        ("2", "Location-Based Organic",       "Kerala spine & ortho clinics, Chennai regenerative medicine centres, Calicut hair loss clinics", "HIGH"),
        ("2", "Content & Authority Sites",    "Practo, Lybrate, Healthline India, OnlyMyHealth — ranking for condition/treatment queries",     "HIGH"),
        ("3", "UAE / NRI Medical Tourism",    "Medical tourism portals targeting Keralites abroad — UAE, Qatar, Kuwait search segment",        "MEDIUM"),
        ("3", "Paid Advertisers",             "Any clinic running Google Ads on BOFU cost/booking keywords in Regencare's clusters",           "MEDIUM"),
        ("3", "AI Answer Citations",          "Clinics or content sites being cited by ChatGPT / Gemini / Perplexity for our target queries", "MEDIUM"),
    ],
    col_widths=[1.2, 4.5, 8.0, 2.3],
)

# ══════════════════════════════════════════════════════════════════════════════
# PHASE 1
# ══════════════════════════════════════════════════════════════════════════════
phase_heading(doc, "PHASE 1", "Competitor Identification",
    "Establish the true SERP competitor list — not who Regencare thinks competes, but who Google shows on the same page.")

sub_heading(doc, "What We Do")
bullets_p1 = [
    "Run every Critical and High-priority keyword from the 142-keyword baseline dataset through Google (incognito, geolocated to Kochi, Calicut, Chennai, and UAE).",
    "Record the top-10 organic results, local pack entries, and featured snippet holder for every keyword.",
    "Tally which domains appear most frequently across all keyword clusters — the most recurring domains are the true SERP competitors.",
    "Separate SERP competitors from business competitors — a Practo listing or a health blog may outrank Regencare but is not a clinic; each requires a different strategic response.",
    "Flag orthogencare.com immediately — document every SERP where it appears alongside regencare.in for the same query (cannibalization evidence file).",
    "Confirm paid competitor presence: note any domains appearing in both organic and paid positions on the same SERP.",
]
for b in bullets_p1: bullet(doc, b)

sub_heading(doc, "Tools")
make_table(doc,
    headers=["Tool", "Purpose", "Cost"],
    rows=[
        ("Google Search (Incognito + Location)", "Manual SERP capture per keyword, per location",             "Free"),
        ("Google Search Console",                "Verify Regencare's own impressions and click data per query","Free"),
        ("Ubersuggest / Ahrefs Free Tier",       "Quick domain overview — top pages, estimated traffic",      "Free tier"),
        ("SEOquake Chrome Extension",            "Instant on-page SEO snapshot while browsing competitor pages","Free"),
        ("Python + BeautifulSoup (custom)",      "Automated SERP scraping for bulk keyword sets",             "Free"),
    ],
    col_widths=[5.5, 9.0, 2.5],
)

sub_heading(doc, "Output")
bullet(doc, "Confirmed competitor list: top-5 SERP competitors per keyword cluster, ranked by frequency of appearance.")
bullet(doc, "Cannibalization evidence file: every query where orthogencare.com and regencare.in co-appear.")

# ══════════════════════════════════════════════════════════════════════════════
# PHASE 2
# ══════════════════════════════════════════════════════════════════════════════
phase_heading(doc, "PHASE 2", "SERP & Ranking Analysis",
    "Map exact positions, identify SERP features owned by competitors, and quantify the ranking gap.")

sub_heading(doc, "What We Do")
bullets_p2 = [
    "For each Critical-priority keyword, record: competitor domain, exact position, page title as shown in SERP, SERP feature held (featured snippet / local pack / PAA / video carousel / image pack).",
    "Identify which SERP features competitors currently hold that Regencare has zero presence in — these are the highest-leverage takeover opportunities.",
    "Map People Also Ask (PAA) boxes: record every PAA question appearing on Regencare's priority SERPs — these are content briefs ready-made.",
    "Identify keywords where the current first-page holder has a weak answer (thin content, no schema, old publish date) — these are the fastest snippet takeover targets.",
    "Track local pack results separately: which clinics appear in the Google Maps 3-pack for 'regenerative medicine Kochi', 'GFC therapy Calicut', and equivalent queries.",
    "Document video carousel presence — keyword clusters where YouTube videos appear on page 1 are priority targets for the video strategy in Step 4B.",
]
for b in bullets_p2: bullet(doc, b)

sub_heading(doc, "Output")
bullet(doc, "SERP position map: table of top-5 competitors per keyword with position, page title, and SERP feature held.")
bullet(doc, "SERP feature gap list: every featured snippet, PAA box, video carousel, and local pack entry we can target.")
bullet(doc, "PAA question bank: full list of related questions appearing on our priority SERPs — fed into content briefs.")

# ══════════════════════════════════════════════════════════════════════════════
# PHASE 3
# ══════════════════════════════════════════════════════════════════════════════
phase_heading(doc, "PHASE 3", "On-Page & Content Audit",
    "Analyse what competitor pages look like structurally and editorially — find the gaps and the white space.")

sub_heading(doc, "What We Do")
bullets_p3 = [
    "For each top-ranking competitor page (one page per keyword cluster): record page title, meta description, H1, H2 structure, word count, number of images, presence of video, schema markup types used, and internal linking pattern.",
    "Assess content depth and quality: does the page actually answer the search query comprehensively, or is it thin? Thin competitor content = easy takeover opportunity.",
    "Identify content formats competitors use: long-form guide, FAQ page, before/after case study, comparison table, clinical explainer — map format to SERP feature won.",
    "Note original differentiators: do competitors use doctor quotes, original research, custom diagrams, clinical data, or patient statistics? If not, these become Regencare's exclusive advantages.",
    "Check NMC compliance gaps in competitor content — clinics making prohibited claims (guaranteed results, comparative superiority) are vulnerable to ranking drops; Regencare builds compliant content that endures.",
    "Map which topics, angles, and question types are covered by every top competitor — then identify the questions nobody has answered well yet (white space content).",
    "Specifically audit orthogencare.com pages that overlap with regencare.in keyword clusters — document exact page-level overlap for the cannibalization resolution strategy.",
]
for b in bullets_p3: bullet(doc, b)

sub_heading(doc, "Output")
bullet(doc, "On-page audit table: structured data for every audited competitor page.")
bullet(doc, "Content gap map: topics, formats, and angles competitors own vs. what nobody owns yet.")
bullet(doc, "White space content list: priority content briefs for topics with no strong current answer in SERPs.")

# ══════════════════════════════════════════════════════════════════════════════
# PHASE 4
# ══════════════════════════════════════════════════════════════════════════════
phase_heading(doc, "PHASE 4", "Keyword Gap Analysis",
    "Find every valuable query competitors rank for that Regencare currently does not.")

sub_heading(doc, "What We Do")
bullets_p4 = [
    "Start with the existing 142-keyword baseline: every 'Not in Top 10' row is the primary gap inventory — 81 keywords identified as absent in the current baseline.",
    "Expand the gap list using Ubersuggest, Ahrefs free, and Google Keyword Planner: pull the top-ranking keywords for each confirmed SERP competitor and cross-reference against our baseline.",
    "Categorise every gap keyword by: intent (informational / commercial / transactional), funnel stage (TOFU / MOFU / BOFU), location cluster (Kochi / Calicut / Chennai / UAE), and estimated search volume.",
    "Prioritise gap keywords by the composite index: search volume x intent value x ranking feasibility divided by competition — same scoring model as the main keyword matrix.",
    "Flag high-CPC gap keywords separately for the Paid Ads launch list — these are priority targets when ad budget is released, even before organic rankings are built.",
    "Identify long-tail gap keywords from the Conversational & Cost cluster — 'how much does GFC therapy cost in Kochi', 'is PRP painful', 'how many sessions does stem cell therapy take' — these are lowest-competition, fastest-rank opportunities.",
    "Add UAE/NRI gap keywords not yet in the baseline: 'regenerative medicine India for NRI', 'Kerala ortho clinic for Keralites in UAE' — low competition, high-value segment.",
]
for b in bullets_p4: bullet(doc, b)

sub_heading(doc, "Output")
bullet(doc, "Keyword gap report: all missing high-value keywords with intent tag, funnel stage, priority score, and recommended target page.")
bullet(doc, "Paid Ads keyword shortlist: BOFU gap keywords ready for immediate campaign activation.")

# ══════════════════════════════════════════════════════════════════════════════
# PHASE 5
# ══════════════════════════════════════════════════════════════════════════════
phase_heading(doc, "PHASE 5", "Backlink Gap Analysis",
    "Identify every high-authority domain linking to competitors but not to Regencare — these become the outreach hit list.")

sub_heading(doc, "What We Do")
bullets_p5 = [
    "Run each top-5 SERP competitor through Ahrefs free backlink checker or Moz Link Explorer — export their top referring domains.",
    "Cross-reference against Regencare.in's own backlink profile — any domain that links to a competitor but not to Regencare is a gap target.",
    "Categorise gap domains by type: Kerala health media, national health portals, medical directories, NRI community sites, university/hospital sites, patient forums.",
    "Score each gap domain by Domain Authority (DA) or Domain Rating (DR) — prioritise outreach to high-DA domains first.",
    "Identify link-earning content opportunities: what type of content earned each backlink for the competitor? (e.g., a clinical guide earned a university citation, a patient story earned a media mention) — replicate the content type.",
    "Build specific outreach angles for each domain type: medical directories get a listing request, health journalists get a story pitch, forums get a doctor-authored answer.",
]
for b in bullets_p5: bullet(doc, b)

sub_heading(doc, "Output")
bullet(doc, "Backlink outreach hit list: domain, DA/DR score, link type, content that earned the link, and recommended outreach angle.")
bullet(doc, "Link-earning content brief: list of content assets that, once published, will naturally attract backlinks from target domains.")

# ══════════════════════════════════════════════════════════════════════════════
# PHASE 6
# ══════════════════════════════════════════════════════════════════════════════
phase_heading(doc, "PHASE 6", "Paid Ads Intelligence",
    "Map the paid advertising landscape — competitor spend patterns, ad copy, landing pages, and offer angles.")

sub_heading(doc, "What We Do")
bullets_p6 = [
    "Run Google searches for all BOFU keywords (cost queries, 'book appointment', 'best clinic' queries) from an incognito browser window geolocated to Kochi, Calicut, and Chennai.",
    "Screenshot and log every Google Ad appearing: headline 1, headline 2, description, display URL, and destination landing page URL.",
    "Visit each competitor landing page linked from an ad: record page structure, primary offer (free consultation, discount, limited slots), CTA copy, trust signals used (certifications, testimonials, case counts).",
    "Repeat searches from a UAE-geolocated session (VPN) for the NRI/medical tourism keyword cluster — competitor paid strategy may differ significantly for that audience.",
    "Use Google's Ads Transparency Centre to cross-check any competitor running sustained campaigns — visible without login.",
    "Identify gaps in competitor ad copy: missing schema extensions, weak CTAs, no price transparency — these are opportunities for Regencare's ads to outperform.",
    "Note which BOFU keywords have zero paid competition — these represent low-cost entry points for Regencare's first ad campaigns.",
]
for b in bullets_p6: bullet(doc, b)

sub_heading(doc, "Output")
bullet(doc, "Paid ads dossier: competitor ad copy, landing page structure, and offer angle per keyword cluster.")
bullet(doc, "Zero-competition BOFU keyword list: high-intent keywords with no current paid competition — priority ad targets.")
bullet(doc, "Landing page best-practice brief: what the best competitor landing pages are doing that Regencare's pages must match or beat.")

# ══════════════════════════════════════════════════════════════════════════════
# PHASE 7
# ══════════════════════════════════════════════════════════════════════════════
phase_heading(doc, "PHASE 7", "Local Search & Google Business Profile Intelligence",
    "Audit the local pack and Google Maps competitive landscape for all three Regencare branch locations.")

sub_heading(doc, "What We Do")
bullets_p7 = [
    "Search 'regenerative medicine [city]', 'GFC therapy [city]', 'hair loss treatment [city]', 'knee pain doctor [city]' for Kochi, Calicut, and Chennai — record all clinics in the local 3-pack.",
    "Visit each competitor's Google Business Profile: category tags, service list, photo count, review count, average rating, Q&A presence, post frequency, and whether booking is enabled.",
    "Compare Regencare's GBP completeness against top local pack competitors — identify missing categories, services, photos, or attributes.",
    "Audit review profile: competitor review volume and recency vs. Regencare across all three branches — a clinic with 200 reviews consistently outranks one with 20, regardless of other factors.",
    "Check NAP (Name, Address, Phone) consistency for each competitor across Google, Practo, JustDial, and Sulekha — if competitors have inconsistencies, note them; we ensure Regencare is cleaner.",
    "Identify which local queries trigger a local pack and which trigger only organic results — local pack presence requires GBP optimisation; organic-only queries require content strategy.",
    "Document the Chennai email/NAP mismatch already identified in the baseline audit — resolve this before any local pack work for Chennai.",
]
for b in bullets_p7: bullet(doc, b)

sub_heading(doc, "Output")
bullet(doc, "Local SEO gap report: side-by-side GBP comparison of Regencare vs. top local pack competitors per city.")
bullet(doc, "GBP improvement checklist: specific missing attributes, categories, photos, and services to add for each Regencare branch.")

# ══════════════════════════════════════════════════════════════════════════════
# PHASE 8
# ══════════════════════════════════════════════════════════════════════════════
phase_heading(doc, "PHASE 8", "AI & Answer Engine Audit (GEO / AEO)",
    "Identify who is being cited by AI engines for Regencare's target queries — and map the path to displacing them.")

sub_heading(doc, "What We Do")
bullets_p8 = [
    "Run 30 priority queries through ChatGPT (GPT-4o), Google Gemini, and Perplexity — record exactly who is cited, named, or linked in the response.",
    "Query types to test: 'best regenerative medicine clinic in Kerala', 'what is GFC therapy India', 'how much does PRP cost in Kochi', 'is stem cell therapy available in India', 'best doctor for knee pain without surgery Kerala'.",
    "Document whether Regencare.in is mentioned in any AI response — and if so, in what context and with what accuracy.",
    "Identify which competitor content is being used as the source for AI answers: the cited URL is the page structure and content format we need to match or improve.",
    "Check Google AI Overviews (SGE): which queries show an AI Overview box? Who is cited in that box? What schema and content format earned that citation?",
    "Identify featured snippet holders on Google for conversational queries — featured snippets are the primary source text for voice answers and often feed AI Overview content.",
    "Record the exact answer format AI engines use for each query type: paragraph (definition queries), list (how-to queries), table (comparison queries) — this informs the answer format library in Step 5.",
]
for b in bullets_p8: bullet(doc, b)

sub_heading(doc, "Output")
bullet(doc, "AI answer competitor map: who is cited per query, what content they used, and what format earned the citation.")
bullet(doc, "GEO gap list: queries where Regencare is entirely absent from AI answers — priority targets for citation-bait content.")
bullet(doc, "Answer format library: paragraph / list / table map per query type for structured content production.")

# ══════════════════════════════════════════════════════════════════════════════
# PHASE 9
# ══════════════════════════════════════════════════════════════════════════════
phase_heading(doc, "PHASE 9", "White Space Identification",
    "Find the topics, formats, and angles that nobody in the SERP owns well — these are Regencare's fastest ranking opportunities.")

sub_heading(doc, "What We Do")
bullets_p9 = [
    "Cross-reference Phase 3 (content gap map), Phase 4 (keyword gap report), and Phase 8 (AI gap list) — the intersection of all three is the white space: high-value queries with no strong current answer anywhere.",
    "Score each white space opportunity: search volume, keyword difficulty, user intent value, and how quickly Regencare could realistically produce the best answer in the SERP.",
    "Identify content formats nobody has used well for the topic — if all competitors have text-only articles on GFC therapy, a video-embedded, schema-marked, FAQ-structured guide is the white space play.",
    "Identify UAE/NRI white space specifically: very few clinics in India have well-optimised content targeting Keralites abroad — 'regenerative medicine in Kerala for NRI', 'medical tourism Kerala ortho' are near-zero-competition.",
    "Flag research and clinical outcome white space: no competitor publishes original data or clinical summaries — Dr. Vineeth MB's experience and outcomes data, published compliantly, becomes a unique authority signal no competitor can copy.",
    "Map white space findings to specific content briefs ready for production in Step 3B.",
]
for b in bullets_p9: bullet(doc, b)

sub_heading(doc, "Output")
bullet(doc, "White space opportunity matrix: topic, format, difficulty score, and recommended content brief for each identified gap.")
bullet(doc, "Quick-win content list: 10 highest-ROI white space topics where Regencare can rank within 60-90 days of publication.")

# ══════════════════════════════════════════════════════════════════════════════
# PHASE 10
# ══════════════════════════════════════════════════════════════════════════════
phase_heading(doc, "PHASE 10", "Positioning Summary",
    "Synthesise all phases into a clear strategic picture: where we win, where we are vulnerable, and where the opening is.")

sub_heading(doc, "What We Do")
bullets_p10 = [
    "Produce a one-page competitive positioning summary for each major keyword cluster (Treatment, Location, Condition, UAE/NRI).",
    "For each cluster, clearly state: WHERE REGENCARE WINS (unique specialist credentials, multi-branch presence, research-backed treatments), WHERE REGENCARE IS VULNERABLE (brand cannibalization, content volume deficit, 404 pages, no booking system), WHERE THE OPENING IS (white space content, UAE segment, AI citation gap).",
    "Produce a consolidated brand cannibalization resolution brief for orthogencare.com: exact strategy recommendation (domain consolidation, 301 redirect, content differentiation, or co-existence strategy with strict keyword separation).",
    "Produce a competitive response brief for GFC therapy: specific actions to move from position 9 to top 3, based on what the 5 clinics currently ranking above are doing and what they are missing.",
    "Present findings in a client-facing summary format: one-page visual matrix showing Regencare vs. top 3 competitors across the 7 key dimensions studied.",
]
for b in bullets_p10: bullet(doc, b)

sub_heading(doc, "Output")
bullet(doc, "Competitive positioning summary: cluster-level win/vulnerable/opportunity table.")
bullet(doc, "Cannibalization resolution brief: specific recommended strategy for orthogencare.com conflict.")
bullet(doc, "GFC therapy battle plan: step-by-step actions to move from #9 to top 3.")
bullet(doc, "Client-facing competitive matrix: visual one-page summary of Regencare vs. key competitors.")

# ─── MASTER DELIVERABLES TABLE ────────────────────────────────────────────────
doc.add_paragraph()
section_banner(doc, "Master Deliverables Summary")
doc.add_paragraph()

make_table(doc,
    headers=["#", "Deliverable", "Fed Into (GTM Step)", "Format"],
    rows=[
        ("D1",  "Confirmed competitor list per cluster",          "Step 2B — Keyword Matrix",       "Excel / Doc"),
        ("D2",  "Cannibalization evidence file (orthogencare.com)","Step 2A — Technical Fixes",      "Doc"),
        ("D3",  "SERP position map",                              "Step 2B — Keyword Matrix",       "Excel"),
        ("D4",  "SERP feature gap list",                          "Step 3A/3B — Content",           "Excel"),
        ("D5",  "PAA question bank",                              "Step 3B — Blog Pipeline",        "Excel"),
        ("D6",  "On-page audit table",                            "Step 3A — Landing Pages",        "Excel"),
        ("D7",  "Content gap map",                                "Step 3B — Blog Pipeline",        "Doc"),
        ("D8",  "White space opportunity matrix",                 "Step 3B — Quick Wins",           "Excel"),
        ("D9",  "Keyword gap report",                             "Step 2B — Keyword Matrix",       "Excel"),
        ("D10", "Backlink outreach hit list",                     "Step 4A — Distribution",         "Excel"),
        ("D11", "Paid ads dossier",                               "Step 2B — Ads Launch List",      "Doc + Screenshots"),
        ("D12", "Local SEO gap report",                           "Step 4A — GBP Optimisation",     "Doc"),
        ("D13", "AI answer competitor map",                       "Step 5C/5D — GEO/AEO",           "Doc"),
        ("D14", "Competitive positioning summary (per cluster)",  "Step 2C — Final Output",         "Doc"),
        ("D15", "Cannibalization resolution brief",               "Step 2A — Priority Fix",         "Doc"),
        ("D16", "GFC therapy battle plan",                        "Step 2B / 3B",                   "Doc"),
        ("D17", "Client-facing competitive matrix (visual)",      "Client Presentation",            "Slide / Doc"),
    ],
    col_widths=[1.0, 6.0, 5.5, 4.5],
)

# ─── TOOLSET TABLE ────────────────────────────────────────────────────────────
doc.add_paragraph()
section_banner(doc, "Complete Toolset")
doc.add_paragraph()

make_table(doc,
    headers=["Tool", "Phase Used", "Purpose", "Cost"],
    rows=[
        ("Google Search (Incognito)",         "P1, P2, P6, P7",  "SERP capture, paid ad capture, local pack audit",     "Free"),
        ("Google Search Console",             "P1, P2",          "Regencare's own impression/click data",               "Free"),
        ("Google Business Profile",           "P7",              "Competitor GBP audit and comparison",                 "Free"),
        ("Google Ads Transparency Centre",    "P6",              "Competitor paid ad history without login",            "Free"),
        ("Ubersuggest (Free Tier)",           "P1, P4",          "Competitor top pages, keyword gaps, domain overview", "Free"),
        ("Ahrefs Free Backlink Checker",      "P5",              "Competitor backlink profile",                         "Free"),
        ("Moz Link Explorer",                 "P5",              "Domain Authority, backlink gap cross-reference",      "Free"),
        ("SEOquake Chrome Extension",         "P2, P3",          "Instant on-page metrics while browsing",              "Free"),
        ("Python + BeautifulSoup",            "P1, P3",          "Automated bulk SERP scraping, on-page data collection","Free"),
        ("Screaming Frog SEO Spider",         "P3",              "Crawl competitor sites for on-page structure data",   "Free (500 URL limit)"),
        ("ChatGPT (GPT-4o)",                  "P8",              "AI answer audit for GEO gap mapping",                 "Free / API"),
        ("Google Gemini",                     "P8",              "AI Overview and Gemini answer audit",                 "Free"),
        ("Perplexity AI",                     "P8",              "AI citation audit — what sources are cited",          "Free"),
        ("Google PageSpeed Insights",         "P2",              "Competitor Core Web Vitals benchmark",                "Free"),
        ("VPN (UAE geolocation)",             "P6, P8",          "Paid ad and AI audit from UAE search context",        "Free / low cost"),
    ],
    col_widths=[4.5, 3.0, 6.5, 3.0],
)

# ─── EXECUTION SEQUENCE ───────────────────────────────────────────────────────
doc.add_paragraph()
section_banner(doc, "Execution Sequence & Parallel Workstreams")
doc.add_paragraph()
body(doc, (
    "Phases 1 through 3 must run sequentially — the competitor list from Phase 1 is the input to Phase 2, "
    "and the Phase 2 SERP map informs which competitor pages to audit in Phase 3. "
    "Phases 4 through 8 can run in parallel once Phase 1 is complete. "
    "Phase 9 (White Space) and Phase 10 (Positioning Summary) are synthesis phases — they require all prior phases to be complete."
), after=8)

make_table(doc,
    headers=["Week", "Phases Running", "Key Dependency", "Output Ready By End of Week"],
    rows=[
        ("Week 1", "P1 (Competitor ID)",                    "142-keyword baseline (complete)",          "Confirmed competitor list"),
        ("Week 1", "P2 (SERP Analysis)",                    "P1 complete",                              "SERP position map + feature gap list"),
        ("Week 2", "P3 (On-Page Audit)",                    "P2 complete — top pages identified",       "On-page audit table + content gap map"),
        ("Week 2", "P4 (Keyword Gap) — parallel",           "P1 complete",                              "Keyword gap report"),
        ("Week 2", "P5 (Backlink Gap) — parallel",          "P1 complete",                              "Backlink outreach hit list"),
        ("Week 2", "P6 (Paid Ads) — parallel",              "P1 complete",                              "Paid ads dossier"),
        ("Week 2", "P7 (Local / GBP) — parallel",           "P1 complete",                              "Local SEO gap report"),
        ("Week 3", "P8 (AI / GEO / AEO Audit)",             "P1 + P3 complete",                         "AI answer competitor map"),
        ("Week 3", "P9 (White Space) — synthesis",          "P3 + P4 + P8 complete",                    "White space opportunity matrix"),
        ("Week 3", "P10 (Positioning Summary) — synthesis", "All phases complete",                      "Full competitive intelligence pack"),
    ],
    col_widths=[1.8, 5.0, 5.0, 5.2],
)

# ─── NEXT ACTION ──────────────────────────────────────────────────────────────
doc.add_paragraph()
section_banner(doc, "NEXT ACTION")
doc.add_paragraph()
body(doc, (
    "Phase 1 (Competitor Identification) begins immediately using the existing 142-keyword baseline dataset. "
    "The first task is the orthogencare.com cannibalization audit — documenting every SERP where it appears alongside regencare.in. "
    "This single finding has the highest immediate ROI of any action in the competitive study."
), after=6)

note_box(doc,
    "PRIORITY REMINDER:",
    (
        "The brand cannibalization issue (orthogencare.com) must be resolved before any content or SEO investment is made. "
        "Every rupee and hour spent building Regencare's rankings is partially undermined as long as a competing domain "
        "with the same doctor is appearing in the same SERPs. Phase 1 and the cannibalization brief are the immediate first output."
    ),
    fill="FFF3CD", label_color=AMB
)

doc.add_paragraph()
footer_p = doc.add_paragraph()
para_fmt(footer_p, WD_ALIGN_PARAGRAPH.LEFT, 12, 2)
pPr = footer_p._p.get_or_add_pPr()
shd = OxmlElement('w:shd'); shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),'F4F6F9')
pPr.append(shd)
add_run(footer_p, "  Prepared by NT Global Digital  ·  Client: Regencare.in  ·  Document Version: 1.0  ·  May 2026  ·  Confidential", size=9, color=NAVY, italic=True)

# ─── SAVE ─────────────────────────────────────────────────────────────────────
out = r"c:\project\AI RESEARCH INTELLIGENCE SYSTEM\Regencare_Competitive_Study_Plan.docx"
doc.save(out)
print(f"Saved: {out}")
